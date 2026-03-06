"""
Engine2: DocxOptimizerEngine - Updated Version
===============================================
Engine untuk merapikan dokumen Word sesuai standar ISO/SNI
Digunakan oleh app.py untuk menu "2. Rapikan (Word -> ISO Std)"
"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_TAB_ALIGNMENT


# Namespace
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
RNS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

def remove_all_hyperlinks(doc):
    """
    Hapus semua hyperlink dalam dokumen dan jadikan teks biasa.
    Mempertahankan formatting run (bold, italic, font size, dll).
    """
    from lxml import etree
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Cari semua elemen hyperlink di seluruh document body
    body = doc.element.body
    hyperlinks = body.findall(f'.//{{{W}}}hyperlink')
    
    for hyperlink in hyperlinks:
        parent = hyperlink.getparent()
        if parent is None:
            continue
        
        # Dapatkan posisi hyperlink di parent
        idx = list(parent).index(hyperlink)
        
        # Pindahkan semua child <w:r> dari hyperlink ke parent (menggantikan hyperlink)
        children = list(hyperlink)
        for i, child in enumerate(children):
            parent.insert(idx + i, child)
        
        # Hapus elemen hyperlink (sudah kosong)
        parent.remove(hyperlink)


def _has_image(paragraph):
    """Cek apakah paragraf mengandung gambar"""
    p_el = paragraph._element
    drawing_tag = f'{{{WNS}}}drawing'
    pict_tag    = f'{{{WNS}}}pict'
    for descendant in p_el.iter():
        if descendant.tag in (drawing_tag, pict_tag):
            return True
    return False

def _is_truly_empty(paragraph):
    """Paragraf kosong = tidak ada teks DAN tidak ada gambar"""
    if _has_image(paragraph):
        return False
    return not paragraph.text.strip()

def set_document_margins(doc, top_cm=3, inside_cm=3, bottom_cm=2, outside_cm=2):
    """Set margin untuk semua sections"""
    for section in doc.sections:
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)
        section.left_margin = Cm(inside_cm)
        section.right_margin = Cm(outside_cm)

#headerfooter
def setup_headers_footers(doc, doc_title="SNI ISO XXXXX:2025", copyright_text="©BSN 2025"):

    # =====================================================
    # HELPER
    # =====================================================
    def clear_container(container):
        container.is_linked_to_previous = False
        for tbl in list(container.tables):
            tbl._element.getparent().remove(tbl._element)
        for p in list(container.paragraphs):
            p._element.getparent().remove(p._element)

    def add_field(run, field_name):
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = field_name

        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)

    # =====================================================
    # NORMALISASI SECTION (A4 + MIRROR)
    # =====================================================
    for section in doc.sections:

        # ---------- FORCE A4 ----------
        top = section.top_margin
        bottom = section.bottom_margin
        left = section.left_margin
        right = section.right_margin
        header_dist = section.header_distance
        footer_dist = section.footer_distance

        section.page_width = Cm(21)
        section.page_height = Cm(29.7)


        section.top_margin = top
        section.bottom_margin = bottom
        section.left_margin = left
        section.right_margin = right
        section.header_distance = header_dist
        section.footer_distance = footer_dist
        section.footer_distance = Pt(35)   # ±1.2 cm (ideal ISO look)


        # ---------- FORCE MIRROR ----------
        sectPr = section._sectPr
        for el in sectPr.findall(qn('w:mirrorMargins')):
            sectPr.remove(el)

        mirror = OxmlElement('w:mirrorMargins')
        sectPr.append(mirror)

        # =====================================================
        # AKTIFKAN ODD/EVEN
        # =====================================================
        section.different_first_page_header_footer = False
        section.odd_and_even_pages_header_footer = True

        # =====================================================
        # RESET PAGE NUMBERING
        # =====================================================
        for el in sectPr.findall(qn('w:pgNumType')):
            sectPr.remove(el)

        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), '1')
        sectPr.append(pgNumType)

        # =====================================================
        # HAPUS HEADER/FOOTER LAMA
        # =====================================================
        clear_container(section.header)
        clear_container(section.footer)
        clear_container(section.first_page_header)
        clear_container(section.first_page_footer)
        clear_container(section.even_page_header)
        clear_container(section.even_page_footer)

        # =====================================================
        # HEADER (TIDAK DIUBAH – SESUAI KODE ANDA)
        # =====================================================
        header = section.header
        p_header = header.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_header = p_header.add_run(doc_title)
        run_header.font.name = "Arial"
        run_header.font.size = Pt(12)
        run_header.bold = True

        even_header = section.even_page_header
        p_even_header = even_header.add_paragraph()
        p_even_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run_even = p_even_header.add_run(doc_title)
        run_even.font.name = "Arial"
        run_even.font.size = Pt(12)
        run_even.bold = True

        # =====================================================
        # HITUNG TENGAH PRESISI
        # =====================================================
        usable_width = int(section.page_width - section.left_margin - section.right_margin)
        center_pos = usable_width // 2

        # =====================================================
        # BUILDER FOOTER
        # =====================================================
        def build_footer(container):

            p = container.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1

            tabs = p.paragraph_format.tab_stops
            tabs.clear_all()
            tabs.add_tab_stop(center_pos, WD_TAB_ALIGNMENT.CENTER)

            # COPYRIGHT
            run_left = p.add_run(copyright_text)
            run_left.font.name = "Arial"
            run_left.font.size = Pt(10)
            run_left.bold = True

            p.add_run("\t")

            # PAGE
            run_page = p.add_run()
            run_page.font.name = "Arial"
            run_page.font.size = Pt(10)
            run_page.bold = True
            add_field(run_page, "PAGE")

            # TEXT
            run_mid = p.add_run(" dari ")
            run_mid.font.name = "Arial"
            run_mid.font.size = Pt(10)
            run_mid.bold = True

            # TOTAL GLOBAL
            run_total = p.add_run()
            run_total.font.name = "Arial"
            run_total.font.size = Pt(10)
            run_total.bold = True
            add_field(run_total, "NUMPAGES")

        # FOOTER GANJIL
        build_footer(section.footer)

        # FOOTER GENAP
        build_footer(section.even_page_footer)




class DocxOptimizerEngine:
    """
    Engine untuk optimasi dokumen Word sesuai standar ISO/SNI
    Compatible dengan app.py interface
    """
    
    def process(self, input_path, output_path, font_name="Arial", font_size=11,
                enable_headers=False, doc_title="", copyright_text="©BSN 2025"):
        """
        Process dokumen Word untuk formatting ISO/SNI
        
        Args:
            input_path: Path input .docx
            output_path: Path output .docx
            font_name: Font name (default: Arial)
            font_size: Font size (default: 11)
            enable_headers: Enable header/footer setup (default: False)
            doc_title: Document title for header
            copyright_text: Copyright text for footer
            
        Returns:
            (success: bool, message: str)
        """
        try:
            doc = Document(input_path)

            # Hapus semua hyperlink → jadikan teks biasa
            remove_all_hyperlinks(doc)

            # Fix ukuran font autonumbering "Annex %1" → 12pt (24 half-points)
            # Label "Annex A" dirender dari numbering lvl rPr, bukan dari run paragraf
            _WNS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # Fix heading style sizes AND fonts so auto-generated numbers match
            _sz_val = str(font_size * 2)  # half-points (11pt = 22)
            _styles_el = doc.styles.element
            for _style in _styles_el.iter(f'{{{_WNS_W}}}style'):
                _styleId = _style.get(f'{{{_WNS_W}}}styleId', '')
                if _styleId.startswith('Heading') and 'Char' not in _styleId:
                    _rPr = _style.find(f'{{{_WNS_W}}}rPr')
                    if _rPr is None:
                        from docx.oxml import OxmlElement as _OxmlElement
                        _rPr = _OxmlElement('w:rPr')
                        _style.append(_rPr)
                    # Fix font size
                    for _sz in _rPr.findall(f'{{{_WNS_W}}}sz'):
                        _sz.set(f'{{{_WNS_W}}}val', _sz_val)
                    for _sz in _rPr.findall(f'{{{_WNS_W}}}szCs'):
                        _sz.set(f'{{{_WNS_W}}}val', _sz_val)
                    # Fix font to Arial - remove old rFonts and replace
                    for _rf in _rPr.findall(f'{{{_WNS_W}}}rFonts'):
                        _rPr.remove(_rf)
                    from lxml import etree as _etree
                    _rFonts_new = _etree.SubElement(_rPr, f'{{{_WNS_W}}}rFonts')
                    _rFonts_new.set(f'{{{_WNS_W}}}ascii', font_name)
                    _rFonts_new.set(f'{{{_WNS_W}}}hAnsi', font_name)
                    _rFonts_new.set(f'{{{_WNS_W}}}cs', font_name)
                    _rPr.insert(0, _rFonts_new)

            # Fix numbering.xml rFonts AND Size
            # LOGIKA BARU: Cek Level Text, jika mengandung 'Annex' -> 12pt, lainnya 11pt
            from lxml import etree as _etree
            _num_part = doc.part.numbering_part
            if _num_part is not None:
                _num_root = _num_part._element
                for _lvl in _num_root.iter(f'{{{_WNS_W}}}lvl'):
                    _rPr = _lvl.find(f'{{{_WNS_W}}}rPr')
                    if _rPr is None:
                        _rPr = _etree.SubElement(_lvl, f'{{{_WNS_W}}}rPr')
                    
                    # 1. Replace rFonts
                    for _rf in _rPr.findall(f'{{{_WNS_W}}}rFonts'):
                        _rPr.remove(_rf)
                    _rFonts_new = _etree.Element(f'{{{_WNS_W}}}rFonts')
                    _rFonts_new.set(f'{{{_WNS_W}}}ascii', font_name)
                    _rFonts_new.set(f'{{{_WNS_W}}}hAnsi', font_name)
                    _rFonts_new.set(f'{{{_WNS_W}}}cs', font_name)
                    _rPr.insert(0, _rFonts_new)

                    # 2. Determine Size based on Level Text (Annex vs Others)
                    _sz_val_current = _sz_val # Default 11pt (22 half-points)
                    _lt = _lvl.find(f'{{{_WNS_W}}}lvlText')
                    if _lt is not None:
                        _lt_val = _lt.get(f'{{{_WNS_W}}}val', '') or ''
                        # Jika lvlText mengandung kata 'annex', paksa ukuran 12pt (24 half-points)
                        if 'annex' in _lt_val.lower():
                            _sz_val_current = '24'
                    
                    # Remove old size
                    for _sz in _rPr.findall(f'{{{_WNS_W}}}sz'):
                        _rPr.remove(_sz)
                    for _szCs in _rPr.findall(f'{{{_WNS_W}}}szCs'):
                        _rPr.remove(_szCs)
                    
                    # Add new size
                    _sz_new = _etree.SubElement(_rPr, f'{{{_WNS_W}}}sz')
                    _sz_new.set(f'{{{_WNS_W}}}val', _sz_val_current)
                    
                    _szCs_new = _etree.SubElement(_rPr, f'{{{_WNS_W}}}szCs')
                    _szCs_new.set(f'{{{_WNS_W}}}val', _sz_val_current)

            def clean_format(paragraph, is_heading=False):
                pf = paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after  = Pt(0)
                pf.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)  # selalu 11pt, termasuk heading/pasal

            def is_subpasal_3(text):
                return bool(re.match(r'^3\.\d+(\s|$)', text))

            # Patterns
            re_split_number = re.compile(r'^(\d[\d\.]*\.?)\s+(.*)')
            re_annex_sub = re.compile(r'^([A-Z]\.\d+[\d\.]*\.?)\s+(.*)')
            re_list_item = re.compile(
                r'^(?:[a-z]\)|[a-z]\.|[A-Z]\)|[A-Z]\.|\([a-z]\)|\([0-9]+\)|[ivxlcdm]+\.|[IVXLCDM]+\.)\s+',
                re.IGNORECASE
            )
            re_copyright = re.compile(r'©\s*ISO.*All\s*rights\s*reserved.*', re.IGNORECASE)
            re_bab = re.compile(r'^(BAB|PASAL|CHAPTER|ARTICLE|SECTION)\s+([IVXLCDM]+|\d+)', re.IGNORECASE)

            # Preprocessing: hapus blank paragraphs
            # Kecuali di bagian tail setelah entri bibliography pertama (preserve spacing asli)
            re_bib_entry = re.compile(r'^\[\d+\]')
            past_bibliography = False
            for p in list(doc.paragraphs):
                if not past_bibliography and re_bib_entry.match(p.text.strip()):
                    past_bibliography = True
                if past_bibliography:
                    continue  # Jaga blank asli di section tail
                if _is_truly_empty(p):
                    try:
                        p._element.getparent().remove(p._element)
                    except:
                        pass

            paragraphs = list(doc.paragraphs)
            tables = list(doc.tables)

            # Tambah enter setelah tabel
            for table in tables:
                try:
                    table_element = table._element
                    parent = table_element.getparent()
                    siblings = list(parent)
                    table_index = siblings.index(table_element)
                    
                    need_blank = True
                    if table_index + 1 < len(siblings):
                        next_el = siblings[table_index + 1]
                        p_tag = f'{{{WNS}}}p'
                        if next_el.tag == p_tag:
                            next_text = ''.join(t.text or '' for t in next_el.iter(f'{{{WNS}}}t'))
                            if not next_text.strip():
                                need_blank = False
                    
                    if need_blank:
                        blank_p = doc.add_paragraph("")
                        clean_format(blank_p)
                        table_element.addnext(blank_p._element)
                except:
                    pass

            # Tambah enter setelah gambar
            for p in list(doc.paragraphs):
                if not _has_image(p):
                    continue
                try:
                    p_element = p._element
                    parent = p_element.getparent()
                    siblings = list(parent)
                    p_index = siblings.index(p_element)
                    
                    need_blank = True
                    if p_index + 1 < len(siblings):
                        next_el = siblings[p_index + 1]
                        p_tag = f'{{{WNS}}}p'
                        if next_el.tag == p_tag:
                            next_text = ''.join(t.text or '' for t in next_el.iter(f'{{{WNS}}}t'))
                            if not next_text.strip():
                                need_blank = False
                    
                    if need_blank:
                        blank_p = doc.add_paragraph("")
                        clean_format(blank_p)
                        p_element.addnext(blank_p._element)
                except:
                    pass

            paragraphs = list(doc.paragraphs)

            # State
            title_processed = False
            in_pasal_3_area = False
            in_bibliography_area = False
            in_annex_area = False
            annex_title_zone = 0  # countdown: paragraf setelah "Annex X" yang harus rapat (informative + judul)

            # Iterasi paragraf
            for p in paragraphs:
                txt = p.text.strip()

                if _has_image(p):
                    clean_format(p)
                    continue

                if not txt:
                    clean_format(p)
                    continue

                if re_copyright.search(txt):
                    try:
                        p._element.getparent().remove(p._element)
                    except:
                        pass
                    continue

                # Deteksi
                match_number = re_split_number.match(txt)
                match_annex_sub = re_annex_sub.match(txt)
                match_bab = re_bab.match(txt)
                has_heading_style = p.style and "Heading" in p.style.name
                has_numbering = (p._element.pPr is not None and p._element.pPr.numPr is not None)
                is_bold_para = any(run.bold for run in p.runs)

                if not is_bold_para and p.style:
                    try:
                        if p.style.font and p.style.font.bold:
                            is_bold_para = True
                    except:
                        pass

                is_list_item = bool(re_list_item.match(txt))
                is_list_item_exception = is_list_item or (is_bold_para and re.match(r'^[A-Z]\s{2,}', txt))
                is_note_exception = txt.lower().startswith('note') or txt.lower().startswith('catatan')
                is_small_font_exception = any(run.font.size and run.font.size.pt == 10 for run in p.runs)
                is_term_definition_exception = p.style and any(
                    term_type in p.style.name for term_type in ['Term', 'Definition']
                )

                is_real_heading = bool(
                    (not is_list_item) and (
                        match_number or match_annex_sub or has_heading_style or has_numbering or match_bab or
                        (p.style and p.style.name.upper() in ('ANNEX', 'ANNEX HEADING')) or
                        (is_bold_para and title_processed and not is_list_item_exception 
                         and not is_note_exception and not is_small_font_exception 
                         and not is_term_definition_exception)
                    )
                )


                # ---- GUARD: ANNEX style selalu diproses sebagai special heading,
                #      tidak boleh jatuh ke blok "Judul Utama" meski title_processed=False ----
                _early_annex = p.style and p.style.name.upper() in ('ANNEX', 'ANNEX HEADING')
                if _early_annex and not title_processed:
                    # Tandai title sudah diproses agar paragraf berikutnya tidak dikira judul
                    title_processed = True

                if not title_processed and not is_real_heading:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    clean_format(p)
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.name = font_name
                        # Paksa warna hitam via XML
                        rPr = run._r.get_or_add_rPr()
                        for color_el in rPr.findall(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color'
                        ):
                            rPr.remove(color_el)
                        color_new = OxmlElement('w:color')
                        color_new.set(qn('w:val'), '000000')
                        rPr.insert(0, color_new)
                    for _ in range(2):
                        blank_after = doc.add_paragraph("")
                        clean_format(blank_after)
                        p._element.addnext(blank_after._element)
                    title_processed = True
                    continue

                # B. HEADING
                if is_real_heading:
                    # Figure/Table
                    _is_fig = bool(re.match(r'^(figure|fig\.?|gambar)\b', txt, re.IGNORECASE))
                    _is_tbl = bool(re.match(r'^(table|tabel)\b', txt, re.IGNORECASE))
                    if _is_fig or _is_tbl:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        clean_format(p)
                        for run in p.runs:
                            run.bold = True
                        blank_after = doc.add_paragraph("")
                        clean_format(blank_after)
                        p._element.addnext(blank_after._element)
                        continue

                    # Special heading
                    _is_biblio_title = p.style and p.style.name == 'Biblio Title'
                    _is_annex_style = p.style and p.style.name.upper() in ('ANNEX', 'ANNEX HEADING')
                    _is_special = _is_biblio_title or _is_annex_style or bool(re.match(
                        r'^(bibliography|bibliografi|annex|lampiran|foreword|kata\s+pengantar|index|indeks)',
                        txt, re.IGNORECASE
                    ))
                    if _is_special:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pf = p.paragraph_format
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)

                        if _is_annex_style:
                            # Gunakan exact line spacing agar Annex A / (informative) / judul mepet
                            from docx.enum.text import WD_LINE_SPACING
                            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                            pf.line_spacing = Pt(14)  # 14pt exact = mepet untuk 12pt font
                            # ANNEX: 12pt bold centered
                            _annex_font_size = 12
                            if not p.runs:
                                r = p.add_run(p.text)
                                r.bold = True
                                r.font.size = Pt(_annex_font_size)
                                r.font.name = font_name
                            else:
                                for run in p.runs:
                                    run.bold = True
                                    run.font.size = Pt(_annex_font_size)
                                    run.font.name = font_name

                            # Hapus <w:br/> ganda sebelum teks judul agar mepet dengan (informative)
                            # Pola: run berisi HANYA <w:br/> tanpa teks → hapus jika run berikutnya punya teks
                            W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            runs_el = p._element.findall(f'{{{W}}}r')
                            for ri, r_el in enumerate(runs_el):
                                r_texts = r_el.findall(f'{{{W}}}t')
                                r_brs = r_el.findall(f'{{{W}}}br')
                                # Run hanya berisi br tanpa teks → kandidat untuk dihapus
                                if r_brs and not r_texts:
                                    # Cek apakah run berikutnya punya teks judul (bukan br saja)
                                    if ri + 1 < len(runs_el):
                                        next_el = runs_el[ri + 1]
                                        next_brs = next_el.findall(f'{{{W}}}br')
                                        next_texts = next_el.findall(f'{{{W}}}t')
                                        # Jika run berikutnya JUGA dimulai dengan br sebelum teks → hapus br ini
                                        if next_brs and next_texts:
                                            # Hapus br dari run berikutnya (br ekstra sebelum judul)
                                            for br in next_brs:
                                                next_el.remove(br)
                        else:
                            # Non-ANNEX special headings (Bibliography, Foreword, dll): 11pt
                            pf.line_spacing = 1.0
                            if not p.runs:
                                r = p.add_run(p.text)
                                r.bold = True
                                r.font.size = Pt(font_size)
                                r.font.name = font_name
                            else:
                                for run in p.runs:
                                    run.bold = True
                                    run.font.size = Pt(font_size)
                                    run.font.name = font_name

                        if re.match(r'^(bibliography|bibliografi)', txt, re.IGNORECASE):
                            in_bibliography_area = True

                        # ANNEX style: satu paragraf berisi (informative) + judul
                        # Tambah 3 blank setelah → total 3 sebelum sub-heading pertama
                        if _is_annex_style or re.match(r'^(annex|lampiran)\b', txt, re.IGNORECASE):
                            in_annex_area = True
                            in_pasal_3_area = False  # Reset agar should_add_enter benar di annex
                            annex_title_zone = 0  # reset, tidak dipakai untuk ANNEX style
                            for _ in range(3):
                                blank_after = doc.add_paragraph("")
                                clean_format(blank_after)
                                p._element.addnext(blank_after._element)
                        else:
                            for _ in range(3):
                                blank_after = doc.add_paragraph("")
                                clean_format(blank_after)
                                p._element.addnext(blank_after._element)
                        continue

                    # Regular heading
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    clean_format(p)
                    for run in p.runs:
                        run.bold = True

                    # Track pasal 3
                    txt_lower = txt.lower()
                    if 'terms and definition' in txt_lower:
                        in_pasal_3_area = True
                    elif match_number:
                        num_str = match_number.group(1).strip('.').strip('(').strip(')')
                        first_char = num_str[0] if num_str else '0'
                        if first_char >= '4':
                            in_pasal_3_area = False
                        elif first_char == '3':
                            in_pasal_3_area = True

                    # Tipe heading
                    is_main_chapter = False
                    if match_bab:
                        is_main_chapter = True
                    elif match_number:
                        num_str = match_number.group(1).strip('.').strip('(').strip(')')
                        is_main_chapter = (num_str.count('.') == 0)
                    elif has_heading_style:
                        is_main_chapter = ("Heading 1" in p.style.name or p.style.name == "Heading")
                    elif has_numbering:
                        try:
                            num_pr = p._element.pPr.numPr
                            if num_pr.ilvl is not None:
                                is_main_chapter = (num_pr.ilvl.val == 0)
                            else:
                                is_main_chapter = True
                        except:
                            is_main_chapter = True

                    is_sub3 = is_subpasal_3(txt)
                    # Annex sub-heading: pola huruf (C.1, A.2, dll) ATAU style 'a2' di dalam annex area
                    is_annex_sub = in_annex_area and (
                        bool(match_annex_sub) or
                        (p.style and p.style.name.lower() in ('a2', 'a3', 'annex sub', 'annex subheading'))
                    )

                    # Spacing sebelum: a2 di annex TIDAK pakai insert_before.
                    # Blank sudah diatur oleh paragraf sebelumnya via should_add_enter + extra_for_annex_sub.
                    # Di area bibliography/tail → jangan tambah blank (spacing sudah preserved dari asli)
                    if not is_annex_sub and not in_bibliography_area:
                        num_enters_before = 1 if (is_main_chapter or is_sub3) else 0
                        for _ in range(num_enters_before):
                            blank = p.insert_paragraph_before("")
                            clean_format(blank)

                    # Format ulang
                    if is_annex_sub:
                        # Format sub-heading annex: left align, bold
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if match_annex_sub:
                            p.text = f"{match_annex_sub.group(1)}    {match_annex_sub.group(2)}"
                        for run in p.runs:
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                    elif match_annex_sub:
                        p.text = f"{match_annex_sub.group(1)}    {match_annex_sub.group(2)}"
                        for run in p.runs:
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                    elif match_number and not match_bab:
                        p.text = f"{match_number.group(1)}    {match_number.group(2)}"
                        for run in p.runs:
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(font_size) # 11pt

                    # Spacing setelah
                    # Di area bibliography/tail → jangan tambah blank
                    if not is_sub3 and not in_bibliography_area:
                        blank_after = doc.add_paragraph("")
                        clean_format(blank_after)
                        p._element.addnext(blank_after._element)

                # C. PARAGRAF BIASA
                else:
                    # Bibliography - MATCH SCREENSHOT FORMAT  
                    is_biblio_entry = p.style and 'Biblio' in p.style.name and 'Title' not in p.style.name
                    bib_match = re.match(r'^\[(\d+)\]\s*(.*)', txt, re.DOTALL)
                    if is_biblio_entry or (in_bibliography_area and bib_match):
                        in_bibliography_area = True
                        
                        # CRITICAL: Remove the style first to avoid inherited formatting
                        try:
                            p.style = 'Normal'
                        except:
                            pass
                        
                        # Apply formatting directly to existing runs
                        first_comma_found = False
                        
                        for run in p.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            run.bold = False
                            
                            run_text = run.text or ''
                            
                            # Number runs and tabs are never italic
                            if run_text.strip() in ['[', ']', '\t'] or run_text.strip().isdigit():
                                run.italic = False
                            elif not first_comma_found:
                                if ',' in run_text:
                                    first_comma_found = True
                                    run.italic = False
                                else:
                                    run.italic = False
                            else:
                                run.italic = True
                        
                        # Set paragraph format with hanging indent
                        pf = p.paragraph_format
                        indent_size = Cm(1.25)
                        pf.left_indent = indent_size
                        pf.first_line_indent = -indent_size
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(6)
                        pf.line_spacing = 1.0
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        # Add tab stop
                        pf.tab_stops.clear_all()
                        pf.tab_stops.add_tab_stop(indent_size, WD_TAB_ALIGNMENT.LEFT)
                        
                        continue

                    is_table_title = bool(re.match(r'^(table|tabel)\b', txt, re.IGNORECASE))
                    is_figure_title = bool(re.match(r'^(figure|fig\.?|gambar)\b', txt, re.IGNORECASE))

                    if is_table_title or is_figure_title:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    clean_format(p)

                    # List item
                    if is_list_item:
                        for run in p.runs:
                            run.bold = False

                    is_bold_non_heading = is_bold_para and not is_real_heading and not is_list_item

                    # Example / Contoh → font size 10, italic/bold sesuai formatting asal
                    if re.match(r'^(EXAMPLE|CONTOH)\b', txt, re.IGNORECASE):
                        for run in p.runs:
                            run.font.size = Pt(10)
                            run.font.name = font_name
                            # italic & bold dibiarkan sesuai aslinya (tidak diubah)

                    # Note
                    if txt.lower().startswith('note') or txt.lower().startswith('catatan'):
                        full_text = p.text
                        for run in p.runs[:]:
                            run._element.getparent().remove(run._element)

                        import re as _re
                        m = _re.match(
                            r'^((?:NOTE|CATATAN)\s*\d*\s*(?:to\s+entry\s*)?:?)\s*',
                            full_text, _re.IGNORECASE
                        )
                        if m:
                            bold_part = m.group(1)
                            normal_part = full_text[m.end():]
                        elif ':' in full_text:
                            parts = full_text.split(':', 1)
                            bold_part = parts[0] + ':'
                            normal_part = parts[1] if len(parts) > 1 else ''
                        else:
                            words = full_text.split(None, 1)
                            bold_part = words[0]
                            normal_part = ' ' + words[1] if len(words) > 1 else ''

                        run_bold = p.add_run(bold_part)
                        run_bold.bold = True
                        run_bold.font.size = Pt(10)
                        run_bold.font.name = font_name

                        if normal_part:
                            run_normal = p.add_run(' ' + normal_part.lstrip())
                            run_normal.bold = False
                            run_normal.font.size = Pt(10)
                            run_normal.font.name = font_name

                    # Spacing setelah
                    should_add_enter = True
                    if in_pasal_3_area:
                        should_add_enter = not is_bold_non_heading
                    elif in_bibliography_area:
                        should_add_enter = False  # Jaga spacing asli di bibliography/tail
                    else:
                        current_idx = paragraphs.index(p)
                        if current_idx + 1 < len(paragraphs):
                            next_p = paragraphs[current_idx + 1]
                            next_txt = next_p.text.strip()
                            if is_subpasal_3(next_txt):
                                should_add_enter = True

                    # Jangan tambah blank di paragraf terakhir
                    current_idx = paragraphs.index(p)
                    is_last_para = (current_idx == len(paragraphs) - 1)

                    if should_add_enter and not is_last_para:
                        blank_after = doc.add_paragraph("")
                        clean_format(blank_after)
                        p._element.addnext(blank_after._element)

                    # Di area annex: tambah 1 blank ekstra jika paragraf berikutnya a2
                    # Rule dari referensi:
                    #   Body Text  → 1 blank → a2
                    #   Note       → 2 blank → a2
                    #   List       → 2 blank → a2  (blank ekstra pakai style List Continue 1)
                    if in_annex_area and should_add_enter and not is_last_para:
                        current_idx = paragraphs.index(p)
                        if current_idx + 1 < len(paragraphs):
                            next_p = paragraphs[current_idx + 1]
                            next_s = next_p.style.name if next_p.style else ''
                            next_is_a2 = next_s.lower() in ('a2', 'a3')
                            cur_s = p.style.name if p.style else ''
                            is_body_text = cur_s in ('Body Text', 'Normal')
                            if next_is_a2 and not is_body_text:
                                extra = doc.add_paragraph("")
                                clean_format(extra)
                                p._element.addnext(extra._element)

            # Set margins
            set_document_margins(doc, top_cm=3, inside_cm=3, bottom_cm=2, outside_cm=2)

            # Setup headers/footers (jika diminta)
            if enable_headers and doc_title:
                setup_headers_footers(doc, doc_title, copyright_text)


            doc.save(output_path)
            return True, output_path

        except Exception as e:
            import traceback
            return False, f"Gagal: {str(e)}\n{traceback.format_exc()}"