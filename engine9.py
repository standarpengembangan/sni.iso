"""
Engine9: DocxFinalTranslatorEngine + 2 Spreadsheet Terpisah (FIXED CORRUPT)
============================================================================
FIX v2:
  - FIX CRITICAL: Hyperlink handling sekarang AMAN, tidak merusak struktur XML
  - 2 SPREADSHEET TERPISAH:
    1. KAMUS_SPREADSHEET_URL → untuk terjemahan istilah
    2. ITALIC_SPREADSHEET_URL → daftar kata yang TIDAK diterjemahkan, OUTPUT MIRING
  - Link Handling: Teks link diterjemahkan, URL diganti placeholder (AMAN).
  - Dual Title Sync, Note/Catatan, Annex fix, dll.
"""

import re
import copy
import time
import uuid
import traceback
import csv
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree


# ─────────────────────────────────────────────────────────────────────────────
# NAMESPACE & KONSTANTA
# ─────────────────────────────────────────────────────────────────────────────

_NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
_W    = f'{{{_NS_W}}}'
_R    = f'{{{_NS_R}}}'

_RE_PURE_NUMBER = re.compile(r'^[\d\s\.\,\:\;\-\(\)\[\]\/\\\+\=\*\%\&\^\$\#\@\!\"\'`~<>{}|_]+$')
_RE_COPYRIGHT = re.compile(r'©|BSN\s*\d{4}', re.IGNORECASE)

_SKIP_STYLES = {
    'caption', 'header', 'footer',
    'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5',
    'table of figures', 'footnote text', 'endnote text', 'macro text',
}
_BIBLIO_TITLE_STYLES = {'biblioti', 'bibliotitle', 'bibliography title'}
_BIBLIO_KEYWORDS_EXACT = {
    'bibliografi', 'bibliography',
    'daftar acuan', 'daftar pustaka', 'daftar referensi',
}
_ANNEX_STYLE_IDS = {'ANNEX', 'Annex', 'annex'}
_HEADING_STYLES_WITH_NUM = {
    'Heading1', 'Heading2', 'Heading3',
    'ANNEX', 'a2', 'a3',
    'Heading4', 'Heading5', 'Heading6',
}
_TRANSLATE_DELAY = 0.15
_EM_DASH = '—'

_LINK_PLACEHOLDER_BASE = "https://placeholder-link.local/"
_LINK_COUNTER = 0

def _get_next_link_placeholder() -> str:
    global _LINK_COUNTER
    _LINK_COUNTER += 1
    return f"{_LINK_PLACEHOLDER_BASE}link-{_LINK_COUNTER}"


# ─────────────────────────────────────────────────────────────────────────────
# 2 URL SPREADSHEET TERPISAH
# ─────────────────────────────────────────────────────────────────────────────

KAMUS_SPREADSHEET_URL = "https://docs.google.com/spreadsheets/d/1BBPCMPwvbBk5LPdoDQwnjQzcPHv7_RDKENqeMsklF-8/edit?usp=sharing"
ITALIC_SPREADSHEET_URL = "https://docs.google.com/spreadsheets/d/1SQnWSA8c1OBVq3XYE8CDMumkt0hpFtxKUjdinfWrqak/edit?usp=sharing"


# ─────────────────────────────────────────────────────────────────────────────
# CUSTOM DICTIONARY (SPREADSHEET 1 - KAMUS TERJEMAHAN)
# ─────────────────────────────────────────────────────────────────────────────

class CustomDictionary:
    """Kamus istilah: source → target (AKAN diterjemahkan)."""
    
    def __init__(self):
        self._entries: dict[str, tuple[str, str]] = {}

    def add_term(self, source: str, target: str) -> None:
        s = source.strip()
        t = target.strip()
        if s and t:
            self._entries[s.lower()] = (s, t)

    def clear(self) -> None:
        self._entries.clear()

    def load_defaults(self) -> int:
        try:
            return self.load_from_google_sheet(KAMUS_SPREADSHEET_URL)
        except Exception:
            return 0

    def load_from_csv(self, filepath: str, src_col: str = 'source', 
                      tgt_col: str = 'target', delimiter: str = ',', 
                      encoding: str = 'utf-8-sig') -> int:
        if not os.path.isfile(filepath): 
            raise FileNotFoundError(f"File CSV tidak ditemukan: {filepath}")
        count = 0
        with open(filepath, newline='', encoding=encoding) as f:
            sample = f.read(1024); f.seek(0)
            has_header = csv.Sniffer().has_header(sample)
            reader = csv.DictReader(f, delimiter=delimiter) if has_header else csv.reader(f, delimiter=delimiter)
            for row in reader:
                if has_header:
                    src = row.get(src_col, row.get('source', '')).strip()
                    tgt = row.get(tgt_col, row.get('target', '')).strip()
                else:
                    row = list(row)
                    if len(row) < 2: continue
                    src, tgt = row[0].strip(), row[1].strip()
                if src and tgt:
                    self._entries[src.lower()] = (src, tgt)
                    count += 1
        return count

    def load_from_excel(self, filepath: str, sheet_name: str | int = 0, 
                        src_col: str = 'source', tgt_col: str = 'target') -> int:
        if not os.path.isfile(filepath): 
            raise FileNotFoundError(f"File Excel tidak ditemukan: {filepath}")
        try: import pandas as pd
        except ImportError: raise ImportError("Jalankan: pip install pandas openpyxl")
        df = pd.read_excel(filepath, sheet_name=sheet_name, dtype=str).fillna('')
        col_src = _find_col(df.columns.tolist(), [src_col, 'source', 'Inggris'])
        col_tgt = _find_col(df.columns.tolist(), [tgt_col, 'target', 'Indonesia'])
        if col_src is None: col_src = df.columns[0]
        if col_tgt is None and len(df.columns) >= 2: col_tgt = df.columns[1]
        if col_tgt is None: raise ValueError("Kolom target tidak ditemukan.")
        count = 0
        for _, row in df.iterrows():
            src = str(row[col_src]).strip()
            tgt = str(row[col_tgt]).strip()
            if src and tgt and src.lower() not in ('nan', '') and tgt.lower() not in ('nan', ''):
                self._entries[src.lower()] = (src, tgt); count += 1
        return count

    def load_from_google_sheet(self, url: str, src_col: str = 'source', 
                               tgt_col: str = 'target', timeout: int = 15) -> int:
        try: import urllib.request, io
        except ImportError: raise ImportError("urllib tidak tersedia.")
        csv_url = _google_sheet_to_csv_url(url)
        try:
            req = urllib.request.Request(csv_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=timeout) as resp: 
                raw = resp.read().decode('utf-8-sig')
        except Exception as e: raise ConnectionError(f"Gagal mengambil data: {e}")
        f = io.StringIO(raw); reader = csv.DictReader(f); fieldnames = reader.fieldnames or []
        col_src = _find_col(fieldnames, [src_col, 'source', 'Inggris'])
        col_tgt = _find_col(fieldnames, [tgt_col, 'target', 'Indonesia'])
        if col_src is None and len(fieldnames) >= 1: col_src = fieldnames[0]
        if col_tgt is None and len(fieldnames) >= 2: col_tgt = fieldnames[1]
        if col_tgt is None: raise ValueError(f"Kolom target tidak ditemukan. Header: {fieldnames}")
        count = 0
        for row in reader:
            src = str(row.get(col_src, '')).strip()
            tgt = str(row.get(col_tgt, '')).strip()
            if src and tgt and src.lower() not in ('', 'nan') and tgt.lower() not in ('', 'nan'):
                self._entries[src.lower()] = (src, tgt); count += 1
        return count

    def __len__(self) -> int: return len(self._entries)
    def list_terms(self) -> list[tuple[str, str]]:
        return [(s, t) for _, (s, t) in sorted(self._entries.items())]

    def _apply_pre(self, text: str) -> tuple[str, dict]:
        if not self._entries: return text, {}
        token_map = {}; result = text.replace(' ', ' ')
        sorted_entries = sorted(self._entries.items(), key=lambda x: len(x[0]), reverse=True)
        for src_lower, (src_orig, tgt) in sorted_entries:
            pattern = re.compile(r'(?<![A-Za-z0-9])' + re.escape(src_lower) + r'(?![A-Za-z0-9])', re.IGNORECASE)
            if pattern.search(result):
                token = f'@@TK_{uuid.uuid4().hex[:8].upper()}@@'
                token_map[token] = tgt
                result = pattern.sub(token, result)
        return result, token_map

    def _apply_post(self, translated: str, token_map: dict) -> str:
        result = translated
        for token, tgt in token_map.items(): 
            result = re.sub(re.escape(token), tgt, result, flags=re.IGNORECASE)
        return result


# ─────────────────────────────────────────────────────────────────────────────
# ITALIC DICTIONARY (SPREADSHEET 2 - KATA MIRING)
# ─────────────────────────────────────────────────────────────────────────────

class ItalicDictionary:
    """
    Daftar kata/frasa dari Spreadsheet 2.
    TIDAK diterjemahkan, otomatis bercetak MIRING di dokumen output.
    """

    def __init__(self):
        self._entries: dict[str, str] = {}

    def add_term(self, term: str) -> None:
        t = term.strip()
        if t:
            self._entries[t.lower()] = t

    def clear(self) -> None:
        self._entries.clear()

    def load_defaults(self) -> int:
        try:
            return self.load_from_google_sheet(ITALIC_SPREADSHEET_URL)
        except Exception:
            return 0

    def load_from_csv(self, filepath: str, term_col: str = 'term', 
                      delimiter: str = ',', encoding: str = 'utf-8-sig') -> int:
        if not os.path.isfile(filepath): 
            raise FileNotFoundError(f"File CSV tidak ditemukan: {filepath}")
        count = 0
        with open(filepath, newline='', encoding=encoding) as f:
            sample = f.read(1024); f.seek(0)
            has_header = csv.Sniffer().has_header(sample)
            reader = csv.DictReader(f, delimiter=delimiter) if has_header else csv.reader(f, delimiter=delimiter)
            skip_vals = {'nan', '', 'term', 'kata', 'italic', 'word', 'text'}
            for row in reader:
                if has_header:
                    term = ''
                    for col_name in [term_col, 'term', 'kata', 'italic', 'word', 'text']:
                        if col_name in row:
                            term = row.get(col_name, '').strip()
                            break
                    if not term and len(row) > 0:
                        term = list(row.values())[0].strip()
                else:
                    row = list(row)
                    if len(row) < 1: continue
                    term = row[0].strip()
                if term and term.lower() not in skip_vals:
                    self._entries[term.lower()] = term; count += 1
        return count

    def load_from_excel(self, filepath: str, sheet_name: str | int = 0, 
                        term_col: str = 'term') -> int:
        if not os.path.isfile(filepath): 
            raise FileNotFoundError(f"File Excel tidak ditemukan: {filepath}")
        try: import pandas as pd
        except ImportError: raise ImportError("Jalankan: pip install pandas openpyxl")
        df = pd.read_excel(filepath, sheet_name=sheet_name, dtype=str).fillna('')
        col_term = None
        for col_name in [term_col, 'term', 'kata', 'italic', 'word', 'text']:
            if col_name in df.columns:
                col_term = col_name; break
        if col_term is None: col_term = df.columns[0]
        count = 0
        skip_vals = {'nan', '', 'term', 'kata', 'italic'}
        for _, row in df.iterrows():
            term = str(row[col_term]).strip()
            if term and term.lower() not in skip_vals:
                self._entries[term.lower()] = term; count += 1
        return count

    def load_from_google_sheet(self, url: str, term_col: str = 'term', 
                               timeout: int = 15) -> int:
        try: import urllib.request, io
        except ImportError: raise ImportError("urllib tidak tersedia.")
        csv_url = _google_sheet_to_csv_url(url)
        try:
            req = urllib.request.Request(csv_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=timeout) as resp: 
                raw = resp.read().decode('utf-8-sig')
        except Exception as e: raise ConnectionError(f"Gagal mengambil data dari Spreadsheet Italic: {e}")
        f = io.StringIO(raw); reader = csv.DictReader(f); fieldnames = reader.fieldnames or []
        col_term = None
        for col_name in [term_col, 'term', 'kata', 'italic', 'word', 'text']:
            if col_name in fieldnames:
                col_term = col_name; break
        if col_term is None and len(fieldnames) >= 1:
            col_term = fieldnames[0]
        if col_term is None: raise ValueError(f"Tidak ada kolom. Header: {fieldnames}")
        count = 0
        skip_vals = {'nan', '', 'term', 'kata', 'italic', 'word', 'text'}
        for row in reader:
            term = str(row.get(col_term, '')).strip()
            if term and term.lower() not in skip_vals:
                self._entries[term.lower()] = term; count += 1
        return count

    def __len__(self) -> int: return len(self._entries)
    
    def list_terms(self) -> list[str]:
        return list(set(self._entries.values()))

    def _apply_pre(self, text: str) -> tuple[str, dict]:
        if not self._entries: return text, {}
        token_map = {}
        result = text
        sorted_entries = sorted(self._entries.items(), key=lambda x: len(x[0]), reverse=True)
        for term_lower, term_orig in sorted_entries:
            pattern = re.compile(r'(?<![A-Za-z0-9])' + re.escape(term_lower) + r'(?![A-Za-z0-9])', re.IGNORECASE)
            if pattern.search(result):
                token = f'@@IT_{uuid.uuid4().hex[:8].upper()}@@'
                token_map[token] = term_orig
                result = pattern.sub(token, result)
        return result, token_map

    def _apply_post(self, translated: str, italic_map: dict) -> tuple[str, list[str]]:
        result = translated
        italic_terms_found = []
        for token, original in italic_map.items():
            if re.search(re.escape(token), result, re.IGNORECASE):
                result = re.sub(re.escape(token), original, result, flags=re.IGNORECASE)
                if original not in italic_terms_found:
                    italic_terms_found.append(original)
        return result, italic_terms_found


# ─────────────────────────────────────────────────────────────────────────────
# SHARED HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _find_col(columns: list, candidates: list) -> str | None:
    for c in candidates:
        if c in columns: return c
    return None

def _google_sheet_to_csv_url(url: str) -> str:
    url = url.strip()
    if 'output=csv' in url or 'format=csv' in url: return url
    if 'docs.google.com/spreadsheets' not in url: return url
    m = re.search(r'/spreadsheets/d/([a-zA-Z0-9_-]+)', url)
    if not m: raise ValueError(f"Tidak dapat mengekstrak Sheet ID: {url}")
    sheet_id = m.group(1)
    gid_match = re.search(r'[#&?]gid=(\d+)', url)
    gid_param = f'&gid={gid_match.group(1)}' if gid_match else ''
    return f'https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv{gid_param}'


# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def _skip_text(text: str) -> bool:
    t = text.strip()
    if len(t) < 3: return True
    if _RE_PURE_NUMBER.fullmatch(t): return True
    if _RE_COPYRIGHT.search(t): return True
    return False

def _skip_paragraph(para, past_bibliography: bool = False) -> bool:
    if past_bibliography: return True
    if not para.text.strip(): return True
    for tag in [f'{_W}drawing', f'{_W}pict']:
        if para._element.find('.//' + tag) is not None: return True
    style_name = (para.style.name or '').lower()
    if any(style_name.startswith(s) for s in _SKIP_STYLES): return True
    return False

def _is_biblio_title_para(para) -> bool:
    style_id = ''
    try:
        pStyle = para._element.find(f'{_W}pPr/{_W}pStyle')
        if pStyle is not None: style_id = pStyle.get(f'{_W}val', '').lower()
    except Exception: pass
    if style_id in _BIBLIO_TITLE_STYLES: return True
    txt = para.text.strip().lower()
    return bool(txt) and not txt[0].isdigit() and txt in _BIBLIO_KEYWORDS_EXACT

def _get_para_style_id(para) -> str:
    pStyle = para._element.find(f'{_W}pPr/{_W}pStyle')
    return pStyle.get(f'{_W}val', '') if pStyle is not None else ''

def _get_style_id(el) -> str:
    pStyle = el.find(f'{_W}pPr/{_W}pStyle')
    return pStyle.get(f'{_W}val', '') if pStyle is not None else 'Normal'

def _has_inline_sectpr(para) -> bool:
    pPr = para._element.find(f'{_W}pPr')
    return pPr is not None and pPr.find(f'{_W}sectPr') is not None

def _all_runs_italic(para) -> bool:
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs: return False
    para_style_italic = False
    try:
        if para.style and para.style.font and para.style.font.italic: para_style_italic = True
    except Exception: pass
    for run in text_runs:
        italic = False
        if run.font.italic is True: italic = True
        if not italic:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{_W}i')
                if i_el is not None:
                    val = i_el.get(f'{_W}val', 'true')
                    if val.lower() not in ('false', '0'): italic = True
        if not italic and para_style_italic:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{_W}i')
                if i_el is not None:
                    val = i_el.get(f'{_W}val', 'true')
                    if val.lower() not in ('false', '0'): italic = True
                else: italic = True
            else: italic = True
        if not italic: return False
    return True


# ─────────────────────────────────────────────────────────────────────────────
# LINK HANDLING (FIXED - AMAN, TIDAK MERUSAK XML)
# ─────────────────────────────────────────────────────────────────────────────

def _has_hyperlinks(para) -> bool:
    """Cek apakah paragraf memiliki hyperlink."""
    return para._element.find(f'{_W}hyperlink') is not None

def _translate_hyperlinks_in_para(para, tr) -> None:
    """
    TERJEMAHKAN teks di dalam hyperlink DAN ganti URL dengan placeholder.
    
    ⚠️ METODE AMAN:
    - TIDAK menghapus element <w:hyperlink>
    - TIDAK membuat element XML baru
    - Hanya mengubah teks di dalam <w:t>
    - Hanya mengubah target_ref di relationship file
    """
    p_el = para._element
    
    for hl in p_el.findall(f'{_W}hyperlink'):
        # 1. Dapatkan semua elemen <w:t> di dalam hyperlink ini
        t_els = hl.findall(f'.//{_W}t')
        hl_text = ''.join(t.text or '' for t in t_els).strip()
        
        if not hl_text:
            continue
        
        # 2. Terjemahkan teks (dengan perlindungan italic)
        text_to_translate = hl_text
        italic_map = {}
        
        if tr.italic_dict and len(tr.italic_dict) > 0:
            text_to_translate, italic_map = tr.italic_dict._apply_pre(text_to_translate)
        
        translated, _ = tr.translate_one(text_to_translate, italic_map)
        
        if italic_map and tr.italic_dict:
            translated, _ = tr.italic_dict._apply_post(translated, italic_map)
        
        # 3. Ganti teks di dalam hyperlink (AMAN: hanya mengubah .text)
        if t_els:
            t_els[0].text = translated
            for t_el in t_els[1:]:
                t_el.text = ''
        
        # 4. Ganti URL di relationship file (AMAN: hanya mengubah atribut)
        r_id = hl.get(f'{_R}id', '')
        if r_id:
            try:
                rels = para.part.rels
                if r_id in rels:
                    rels[r_id]._target = _get_next_link_placeholder()
            except Exception:
                pass  # Jika gagal, biarkan URL lama (lebih baik daripada corrupt)


# ─────────────────────────────────────────────────────────────────────────────
# ITALIC FORMATTING
# ─────────────────────────────────────────────────────────────────────────────

def _apply_mixed_formatting_to_para(para, text: str, italic_terms: list[str], 
                                     font_name: str = None, font_size: int = None) -> None:
    """Terapkan formatting: kata di italic_terms jadi MIRING."""
    if not italic_terms:
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]: r.text = ''
        else:
            para.add_run(text)
        return
    
    italic_positions = []
    for term in italic_terms:
        start = 0
        term_lower = term.lower()
        text_lower = text.lower()
        while True:
            idx = text_lower.find(term_lower, start)
            if idx == -1: break
            italic_positions.append((idx, idx + len(term)))
            start = idx + 1
    
    italic_positions.sort(key=lambda x: x[0])
    
    filtered_positions = []
    last_end = -1
    for start, end in italic_positions:
        if start >= last_end:
            filtered_positions.append((start, end))
            last_end = end
    
    if not filtered_positions:
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]: r.text = ''
        else:
            para.add_run(text)
        return
    
    segments = []
    last_pos = 0
    
    for start, end in filtered_positions:
        if start > last_pos:
            segments.append((text[last_pos:start], False))
        segments.append((text[start:end], True))
        last_pos = end
    
    if last_pos < len(text):
        segments.append((text[last_pos:], False))
    
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    
    for seg_text, is_italic in segments:
        if not seg_text: continue
        
        run = para.add_run(seg_text)
        run.font.name = font_name or 'Arial'
        if font_size:
            run.font.size = Pt(font_size)
        
        if is_italic:
            run.italic = True


# ─────────────────────────────────────────────────────────────────────────────
# FITUR 1: REKONSTRUKSI ANNEX
# ─────────────────────────────────────────────────────────────────────────────

def _fix_annex_style_para(para, annex_letter: str = None) -> None:
    sid = _get_para_style_id(para)
    if sid not in _ANNEX_STYLE_IDS: return
    full_text = ''.join(r.text for r in para.runs if r.text is not None).strip()
    if not full_text: return
    tag_norm = None; title_part = full_text
    for t in ['(informatif)', '(normatif)', '(informative)', '(normative)', '(informasi)']:
        idx = full_text.lower().find(t)
        if idx != -1:
            tag_norm = '(normatif)' if 'norm' in t.lower() else '(informatif)'
            title_part = full_text[:idx] + " " + full_text[idx + len(t):]
            break
    annex_label = None
    m_annex = re.match(r'^(?:Annex|Lampiran)\s+([A-Z0-9\.]+)\s*', title_part, flags=re.IGNORECASE)
    if m_annex:
        annex_label = f'Lampiran {m_annex.group(1).upper()}'
        title_part = title_part[m_annex.end():].strip()
    elif annex_letter:
        annex_label = f'Lampiran {annex_letter.upper()}'
        title_part = re.sub(r'^(?:Annex|Lampiran)\s*\S*\s*', '', title_part, flags=re.IGNORECASE).strip()
    else:
        title_part = re.sub(r'^(?:Annex|Lampiran)\s*\S*\s*', '', title_part, flags=re.IGNORECASE).strip()
    title_part = title_part.lstrip('\n').strip()
    pPr = para._element.find(f'{_W}pPr')
    if pPr is None: pPr = etree.SubElement(para._element, f'{_W}pPr')
    old_numPr = pPr.find(f'{_W}numPr')
    if old_numPr is not None: pPr.remove(old_numPr)
    pPr.insert(0, etree.fromstring(f'<w:numPr xmlns:w="{_NS_W}"><w:ilvl w:val="0"/><w:numId w:val="0"/></w:numPr>'))
    for child in list(para._element):
        if child is not pPr: para._element.remove(child)
    def mk(text, bold=False):
        esc = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        b = '<w:b/><w:bCs/>' if bold else ''
        return etree.fromstring(f'<w:r xmlns:w="{_NS_W}"><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>{b}<w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:t xml:space="preserve">{esc}</w:t></w:r>')
    new_runs = []
    if annex_label: new_runs.extend([mk(annex_label, True), etree.fromstring(f'<w:r xmlns:w="{_NS_W}"><w:br/></w:r>')])
    else: new_runs.append(etree.fromstring(f'<w:r xmlns:w="{_NS_W}"><w:br/></w:r>'))
    if tag_norm: new_runs.append(mk(tag_norm, True))
    if title_part: new_runs.extend([etree.fromstring(f'<w:r xmlns:w="{_NS_W}"><w:br/></w:r>'), mk(title_part, True)])
    for run_el in new_runs: para._element.append(run_el)


# ─────────────────────────────────────────────────────────────────────────────
# FITUR 2: EM DASH TO BULLETS
# ─────────────────────────────────────────────────────────────────────────────

def _get_or_create_emdash_numid(doc: Document) -> str:
    try:
        np = doc.part.numbering_part
        if np is None: return None
        nxml = np._element; target_abstract_id = None
        for ab in nxml.findall(f'{_W}abstractNum'):
            for lvl in ab.findall(f'{_W}lvl'):
                txt_el = lvl.find(f'{_W}lvlText')
                if txt_el is not None and txt_el.get(f'{_W}val') == _EM_DASH:
                    target_abstract_id = ab.get(f'{_W}abstractNumId'); break
            if target_abstract_id: break
        if target_abstract_id:
            existing_nums = nxml.findall(f'{_W}num')
            max_num_id = max((int(n.get(f'{_W}numId', 0)) for n in existing_nums), default=0)
            new_num_id = str(max_num_id + 1)
            nxml.append(etree.fromstring(f'<w:num xmlns:w="{_NS_W}" w:numId="{new_num_id}"><w:abstractNumId w:val="{target_abstract_id}"/></w:num>'))
            return new_num_id
        existing_abstracts = nxml.findall(f'{_W}abstractNum')
        max_abstract_id = max((int(a.get(f'{_W}abstractNumId', 0)) for a in existing_abstracts), default=0)
        new_abstract_id = str(max_abstract_id + 1)
        existing_nums = nxml.findall(f'{_W}num')
        max_num_id = max((int(n.get(f'{_W}numId', 0)) for n in existing_nums), default=0)
        new_num_id = str(max_num_id + 1)
        nxml.insert(0, etree.fromstring(f'<w:abstractNum xmlns:w="{_NS_W}" w:abstractNumId="{new_abstract_id}"><w:multiLevelType w:val="hybridMultilevel"/><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="{_EM_DASH}"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr><w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr></w:lvl></w:abstractNum>'))
        nxml.append(etree.fromstring(f'<w:num xmlns:w="{_NS_W}" w:numId="{new_num_id}"><w:abstractNumId w:val="{new_abstract_id}"/></w:num>'))
        return new_num_id
    except Exception as e: print(f"[Engine9] Error numbering: {e}"); return None

def _convert_emdash_to_bullets(doc: Document) -> None:
    num_id = _get_or_create_emdash_numid(doc)
    if not num_id: return
    for para in doc.paragraphs:
        sid = _get_para_style_id(para)
        if sid in _ANNEX_STYLE_IDS or sid.startswith('Heading'): continue
        text = para.text.strip()
        if text.startswith(_EM_DASH):
            for r in para.runs:
                if _EM_DASH in r.text: r.text = r.text.replace(_EM_DASH, "", 1).lstrip(); break
            pPr = para._element.find(f'{_W}pPr')
            if pPr is None: pPr = etree.SubElement(para._element, f'{_W}pPr')
            old_num = pPr.find(f'{_W}numPr')
            if old_num is not None: pPr.remove(old_num)
            pPr.insert(0, etree.fromstring(f'<w:numPr xmlns:w="{_NS_W}"><w:ilvl w:val="0"/><w:numId w:val="{num_id}"/></w:numPr>'))
            old_ind = pPr.find(f'{_W}ind')
            if old_ind is not None: pPr.remove(old_ind)
            pPr.insert(1, etree.fromstring(f'<w:ind xmlns:w="{_NS_W}" w:left="360" w:hanging="360"/>'))


# ─────────────────────────────────────────────────────────────────────────────
# FITUR 3: FIX NOTE / CATATAN
# ─────────────────────────────────────────────────────────────────────────────

def _fix_note_para(para) -> None:
    full_text = para.text.strip().replace(' ', ' ')
    if not full_text: return
    txt_lower = full_text.lower()
    if not (txt_lower.startswith('note') or txt_lower.startswith('catatan')): return
    _has_extra = bool(re.search(r'(?:to\s+entry|untuk\s+entri|untuk\s+masuk|untuk\s+dimasukkan|untuk\s+diterapkan)', full_text, re.IGNORECASE))
    if txt_lower.startswith('note'):
        m = re.match(r'^(NOTE|Note|note)', full_text)
        if m: full_text = ('Catatan' if _has_extra else 'CATATAN') + full_text[len(m.group(1)):]
    elif txt_lower.startswith('catatan'):
        m = re.match(r'^(CATATAN|Catatan|catatan)', full_text)
        if m: full_text = ('Catatan' if _has_extra else 'CATATAN') + full_text[len(m.group(1)):]
    full_text = re.sub(r'((?:CATATAN|Catatan)(?:\s+\d+)?\s+)(untuk\s+masuk|untuk\s+dimasukkan|untuk\s+diterapkan)', lambda mo: mo.group(1) + 'untuk entri', full_text)
    m_colon = re.match(r'^((?:NOTE|CATATAN|Catatan)[^:]*:)\s*', full_text, re.IGNORECASE)
    if m_colon: bold_part, normal_part = m_colon.group(1).rstrip(), full_text[m_colon.end():]
    else:
        words = full_text.split(None, 1); bold_part = words[0]; normal_part = words[1] if len(words) > 1 else ''
    if not normal_part.strip(): return
    font_name = 'Arial'; font_size = None
    for run in para.runs:
        if run.text.strip():
            if run.font.name: font_name = run.font.name
            if run.font.size: font_size = run.font.size
            break
    for run in list(para.runs): run._element.getparent().remove(run._element)
    run_b = para.add_run(bold_part); run_b.bold = True; run_b.font.name = font_name
    if font_size: run_b.font.size = font_size
    run_n = para.add_run(' ' + normal_part.lstrip()); run_n.font.name = font_name
    if font_size: run_n.font.size = font_size

def _fix_all_notes(doc: Document) -> None:
    for para in doc.paragraphs: _fix_note_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs: _fix_note_para(para)


# ─────────────────────────────────────────────────────────────────────────────
# BIBLIOGRAFI & AUTONUMBERING
# ─────────────────────────────────────────────────────────────────────────────

def _el_text(el) -> str: return ''.join(t.text or '' for t in el.findall(f'.//{_W}t')).strip()
def _is_bibliography_el(el) -> bool:
    if el.tag != f'{_W}p': return False
    sid = _get_style_id(el).lower()
    if sid in _BIBLIO_TITLE_STYLES: return True
    text = _el_text(el).lower().strip()
    if not text or text[0].isdigit(): return False
    return text in _BIBLIO_KEYWORDS_EXACT

def _find_bib_index(body_els: list) -> int:
    for i, el in enumerate(body_els):
        if _is_bibliography_el(el): return i
    return -1

def _find_content_start_before_bib(body_els: list, bib_idx: int) -> int:
    search_limit = bib_idx if bib_idx >= 0 else len(body_els)
    last_sectpr = -1
    for i, el in enumerate(body_els):
        if i >= search_limit: break
        if el.tag != f'{_W}p': continue
        pPr = el.find(f'{_W}pPr')
        if pPr is not None and pPr.find(f'{_W}sectPr') is not None: last_sectpr = i
    return last_sectpr + 1

def _read_style_numpr(doc: Document) -> dict:
    result = {}
    try:
        for style_el in doc.part.styles._element.findall(f'{_W}style'):
            sid = style_el.get(f'{_W}styleId', '')
            if sid not in _HEADING_STYLES_WITH_NUM: continue
            pPr = style_el.find(f'{_W}pPr')
            if pPr is None: continue
            numPr = pPr.find(f'{_W}numPr')
            if numPr is None: continue
            nid = numPr.find(f'{_W}numId'); ilvl = numPr.find(f'{_W}ilvl')
            result[sid] = {'numId': nid.get(f'{_W}val', '0') if nid is not None else '0',
                           'ilvl': ilvl.get(f'{_W}val', '0') if ilvl is not None else '0'}
    except Exception: pass
    return result

def _create_restart_numids(doc: Document, style_numpr: dict) -> dict:
    remapping = {}
    if not style_numpr: return remapping
    try:
        nxml = doc.part.numbering_part._element
        unique_numids = {info['numId'] for info in style_numpr.values()}
        existing = nxml.findall(f'{_W}num')
        max_id = max((int(n.get(f'{_W}numId', 0)) for n in existing), default=0)
        for old_nid in unique_numids:
            if old_nid == '0': continue
            old_num_el = None
            for n in existing:
                if n.get(f'{_W}numId') == old_nid: old_num_el = n; break
            if old_num_el is None: continue
            max_id += 1; new_nid = str(max_id)
            new_num_el = copy.deepcopy(old_num_el); new_num_el.set(f'{_W}numId', new_nid)
            for lo in new_num_el.findall(f'{_W}lvlOverride'): new_num_el.remove(lo)
            new_num_el.append(etree.fromstring(f'<w:lvlOverride xmlns:w="{_NS_W}" w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>'))
            nxml.append(new_num_el); remapping[old_nid] = new_nid
    except Exception: pass
    return remapping

def _apply_numpr_restart_to_headings(elements: list, style_numpr: dict, remapping: dict) -> None:
    if not remapping: return
    for el in elements:
        if el.tag != f'{_W}p': continue
        sid = _get_style_id(el)
        if sid not in style_numpr: continue
        new_nid = remapping.get(style_numpr[sid]['numId'])
        if not new_nid: continue
        pPr = el.find(f'{_W}pPr')
        if pPr is None: pPr = etree.SubElement(el, f'{_W}pPr'); el.insert(0, pPr)
        old_numPr = pPr.find(f'{_W}numPr')
        if old_numPr is not None: pPr.remove(old_numPr)
        pPr.insert(0, etree.fromstring(f'<w:numPr xmlns:w="{_NS_W}"><w:ilvl w:val="{style_numpr[sid]["ilvl"]}"/><w:numId w:val="{new_nid}"/></w:numPr>'))


# ─────────────────────────────────────────────────────────────────────────────
# INSERT ORIGINAL
# ─────────────────────────────────────────────────────────────────────────────

def _page_break_para() -> etree._Element:
    return etree.fromstring(f'<w:p xmlns:w="{_NS_W}"><w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr><w:r><w:br w:type="page"/></w:r></w:p>')

def _intro_heading() -> etree._Element:
    return etree.fromstring(f'<w:p xmlns:w="{_NS_W}"><w:pPr><w:jc w:val="center"/><w:spacing w:before="0" w:after="0"/></w:pPr><w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/><w:b/><w:bCs/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:t>Introduction</w:t></w:r></w:p>')

def _empty_para() -> etree._Element:
    return etree.fromstring(f'<w:p xmlns:w="{_NS_W}"><w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr></w:p>')

def _insert_original_before_bib(translated_doc: Document, orig_body_els: list, progress_callback=None) -> None:
    bib_in_orig = _find_bib_index(orig_body_els)
    content_start = _find_content_start_before_bib(orig_body_els, bib_in_orig)
    content_end = bib_in_orig if bib_in_orig >= 0 else len(orig_body_els)
    if bib_in_orig < 0 and content_end > content_start:
        if orig_body_els[content_end - 1].tag == f'{_W}sectPr': content_end -= 1
    els_to_insert = orig_body_els[content_start:content_end]
    if not els_to_insert:
        _notify(progress_callback, 50, "⚠️ Tidak ada konten asli."); return
    _notify(progress_callback, 10, f"Ditemukan {len(els_to_insert)} elemen.")
    style_numpr = _read_style_numpr(translated_doc)
    remapping = _create_restart_numids(translated_doc, style_numpr)
    new_content_els = [copy.deepcopy(el) for el in els_to_insert]
    _apply_numpr_restart_to_headings(new_content_els, style_numpr, remapping)
    _notify(progress_callback, 40, "Reset heading numbers.")
    trans_body = translated_doc.element.body; trans_children = list(trans_body)
    bib_idx_trans = _find_bib_index(trans_children)
    if bib_idx_trans < 0:
        bib_idx_trans = len(trans_children)
        for i in range(len(trans_children) - 1, -1, -1):
            if trans_children[i].tag == f'{_W}sectPr': bib_idx_trans = i; break
    new_els = [_page_break_para(), _intro_heading(), _page_break_para()] + new_content_els + [_empty_para()]
    total = len(new_els)
    for offset, el in enumerate(new_els):
        trans_body.insert(bib_idx_trans + offset, el)
        if progress_callback and offset % 20 == 0:
            _notify(progress_callback, 45 + int(offset / max(total, 1) * 55), f"Menyisipkan... ({offset}/{total})")
    _notify(progress_callback, 100, f"✅ Selesai ({len(els_to_insert)} item).")

def _notify(cb, pct: int, msg: str) -> None:
    if cb:
        try: cb(pct, msg)
        except Exception: pass


# ─────────────────────────────────────────────────────────────────────────────
# TRANSLATOR WRAPPER
# ─────────────────────────────────────────────────────────────────────────────

class _Translator:
    def __init__(self, source: str = 'auto', target: str = 'id', 
                 custom_dict: CustomDictionary | None = None,
                 italic_dict: ItalicDictionary | None = None):
        try: from deep_translator import GoogleTranslator
        except ImportError: raise ImportError("Jalankan: pip install deep-translator")
        self._cls = GoogleTranslator
        self.source = source; self.target = target
        self.custom_dict = custom_dict; self.italic_dict = italic_dict

    def translate_one(self, text: str, italic_map: dict = None) -> tuple[str, list[str]]:
        t = text.strip()
        if not t or _skip_text(t): return text, []
        token_map = {}
        if self.custom_dict and len(self.custom_dict) > 0:
            t, token_map = self.custom_dict._apply_pre(t)
        final_italic_map = italic_map or {}
        try:
            result = self._cls(source=self.source, target=self.target).translate(t)
            if not result: result = t
        except Exception:
            time.sleep(0.8)
            try:
                result = self._cls(source=self.source, target=self.target).translate(t)
                if not result: result = t
            except Exception: result = t
        if token_map: result = self.custom_dict._apply_post(result, token_map)
        italic_terms_found = []
        if final_italic_map and self.italic_dict:
            result, italic_terms_found = self.italic_dict._apply_post(result, final_italic_map)
        return result, italic_terms_found

def _match_capitalization(original: str, translated: str) -> str:
    orig = original.strip(); tran = translated.strip()
    if not orig or not tran: return translated
    letters = [c for c in orig if c.isalpha()]
    if not letters: return translated
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    if upper_ratio >= 0.8: return tran.upper()
    if orig[0].isupper(): return tran[0].upper() + tran[1:] if len(tran) > 1 else tran.upper()
    return tran


def _translate_para(para, tr, past_bibliography: bool = False) -> list[str]:
    """
    Terjemahkan paragraf.
    Hyperlink ditangani SECARA TERPISAH dan AMAN (tidak merusak XML).
    """
    if _skip_paragraph(para, past_bibliography): return []
    
    p_el = para._element
    
    # 1. Cek apakah ada hyperlink
    has_hl = _has_hyperlinks(para)
    
    # 2. Proses TEKS NORMAL (bukan hyperlink)
    text_runs = [(i, r) for i, r in enumerate(para.runs) if r.text and r.text.strip()]
    
    if not text_runs:
        # Jika tidak ada teks normal, hanya terjemahkan hyperlink (jika ada)
        if has_hl:
            _translate_hyperlinks_in_para(para, tr)
        return []
    
    combined = ''.join(r.text for _, r in text_runs).strip()
    if _skip_text(combined): 
        if has_hl:
            _translate_hyperlinks_in_para(para, tr)
        return []
    
    # Get font
    font_name = 'Arial'; font_size = None
    for run in para.runs:
        if run.text.strip():
            if run.font.name: font_name = run.font.name
            if run.font.size: font_size = run.font.size.pt if run.font.size else None
            break
    
    # Proteksi italic
    italic_map = {}
    if tr.italic_dict and len(tr.italic_dict) > 0:
        combined, italic_map = tr.italic_dict._apply_pre(combined)
    
    # Terjemahkan
    translated, italic_terms_found = tr.translate_one(combined, italic_map)
    time.sleep(_TRANSLATE_DELAY)
    if not translated or translated == combined: 
        if has_hl:
            _translate_hyperlinks_in_para(para, tr)
        return []
    translated = _match_capitalization(combined, translated)
    
    # Apply formatting ke teks normal
    if italic_terms_found:
        _apply_mixed_formatting_to_para(para, translated, italic_terms_found, font_name, font_size)
    else:
        if para.runs:
            para.runs[0].text = translated
            for r in para.runs[1:]: r.text = ''
        else:
            para.add_run(translated)
    
    # 3. TERJEMAHKAN HYPERLINK SECARA TERPISAH (AMAN)
    if has_hl:
        _translate_hyperlinks_in_para(para, tr)
    
    return italic_terms_found


def _translate_table(table, tr) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs: _translate_para(para, tr)

def _translate_hf(hf_part, tr) -> None:
    if hf_part is None: return
    try:
        for para in hf_part.paragraphs:
            if _RE_COPYRIGHT.search(para.text or ''): continue
            _translate_para(para, tr)
        for table in hf_part.tables: _translate_table(table, tr)
    except Exception: pass


# ─────────────────────────────────────────────────────────────────────────────
# SINKRONISASI JUDUL
# ─────────────────────────────────────────────────────────────────────────────

def _extract_cover_titles(doc: Document) -> tuple[str, str]:
    id_title = ""; en_title = ""
    for para in doc.paragraphs:
        if _has_inline_sectpr(para): break
        text = para.text.strip()
        if not text: continue
        is_bold = False; is_italic = False; max_size = 0
        for run in para.runs:
            if run.text.strip():
                if run.font.bold: is_bold = True
                if run.font.italic: is_italic = True
                sz = run.font.size
                if sz and sz.pt > max_size: max_size = sz.pt
        if not id_title:
            if max_size >= 16 and is_bold and not is_italic: id_title = text; continue
        if id_title and not en_title:
            if max_size >= 14 and is_italic: en_title = text
    return id_title, en_title

def _sync_body_title(doc: Document, cover_id: str) -> bool:
    if not cover_id: return False
    _BODY_TITLE_STYLES = {'main title 2', 'main title2', 'maintitle2', 'boxedtitle', 'boxed title', 'body title'}
    _RE_H1 = re.compile(r'^\d+\s+\S')
    paras = doc.paragraphs; sect_idx = -1; h1_idx = -1
    for i, p in enumerate(paras):
        if sect_idx == -1 and _has_inline_sectpr(p): sect_idx = i
        txt = p.text.strip()
        if txt and _RE_H1.match(txt) and 'heading' in (p.style.name or '').lower(): h1_idx = i; break
    start = sect_idx + 1 if sect_idx != -1 else 0
    end = h1_idx if h1_idx != -1 else start + 20
    if end > len(paras): end = len(paras)
    for i in range(start, end):
        if (paras[i].style.name or '').lower().strip() in _BODY_TITLE_STYLES and paras[i].text.strip():
            _replace_para_text(paras[i], cover_id); return True
    for i in range(start, end):
        txt = paras[i].text.strip()
        if txt and any(r.bold for r in paras[i].runs if r.text.strip()) and paras[i].alignment == WD_ALIGN_PARAGRAPH.CENTER:
            _replace_para_text(paras[i], cover_id); return True
    return False

def _replace_para_text(para, new_text: str) -> None:
    fn = 'Arial'; fs = None
    for r in para.runs:
        if r.text.strip():
            if r.font.name: fn = r.font.name
            if r.font.size: fs = r.font.size
            break
    for r in list(para.runs): r._element.getparent().remove(r._element)
    nr = para.add_run(new_text); nr.bold = True; nr.font.name = fn
    if fs: nr.font.size = fs
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _sync_foreword_title(doc: Document, cover_id: str, cover_en: str) -> None:
    if not cover_id: return
    pat = re.compile(r'^(SNI [^,]+,\s*)(.*?)(,\s*merupakan standar.*?identik dari [^,]+,\s*)(.*?)(,\s*dengan metode.*)', re.IGNORECASE | re.DOTALL)
    for para in doc.paragraphs:
        m = pat.match(para.text.strip())
        if m:
            if para._element.pPr is None: continue
            for r in list(para.runs): r._element.getparent().remove(r._element)
            def mk(t, ital=False):
                r = para.add_run(t); r.font.name = "Arial"; r.font.size = Pt(11); r.italic = ital; return r
            mk(m.group(1)); mk(cover_id, True); mk(m.group(3))
            if cover_en: mk(cover_en, True)
            mk(m.group(5)); para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; break


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENGINE (FIXED)
# ─────────────────────────────────────────────────────────────────────────────

class DocxFinalTranslatorEngine:
    """
    Engine utama dengan 2 SPREADSHEET TERPISAH.
    FIXED: Hyperlink handling sekarang AMAN (tidak corrupt file).
    """
    
    def __init__(self, source_lang: str = 'auto', target_lang: str = 'id', 
                 custom_dict: CustomDictionary | None = None,
                 italic_dict: ItalicDictionary | None = None):
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.custom_dict = custom_dict
        self.italic_dict = italic_dict

    def set_dictionary(self, d: CustomDictionary) -> None: self.custom_dict = d
    def set_italic_dictionary(self, d: ItalicDictionary) -> None: self.italic_dict = d
    def get_dictionary(self) -> CustomDictionary:
        if self.custom_dict is None: self.custom_dict = CustomDictionary()
        return self.custom_dict
    def get_italic_dictionary(self) -> ItalicDictionary:
        if self.italic_dict is None: self.italic_dict = ItalicDictionary()
        return self.italic_dict

    def translate(self, input_docx: str, output_docx: str, progress_callback=None, 
                  translate_headers: bool = False) -> tuple[bool, str]:
        try:
            global _LINK_COUNTER; _LINK_COUNTER = 0
            
            info = f"{len(self.custom_dict) if self.custom_dict else 0} kamus"
            if self.italic_dict: info += f", {len(self.italic_dict)} miring"
            _notify(progress_callback, 2, f"Snapshot... ({info})")
            
            doc_orig = Document(input_docx)
            orig_body_els = [copy.deepcopy(el) for el in doc_orig.element.body]
            del doc_orig

            _notify(progress_callback, 5, "Init translator...")
            tr = _Translator(self.source_lang, self.target_lang, self.custom_dict, self.italic_dict)
            doc = Document(input_docx)

            body = doc.element.body
            para_map = {p._element: p for p in doc.paragraphs}
            tbl_map = {t._element: t for t in doc.tables}

            items = []; sec_brks = 0; COVER_END = 1
            for child in body:
                if child in para_map:
                    in_cover = (sec_brks < COVER_END)
                    items.append(('para', para_map[child], in_cover))
                    if _has_inline_sectpr(para_map[child]): sec_brks += 1
                elif child in tbl_map:
                    items.append(('table', tbl_map[child], sec_brks < COVER_END))

            total = len(items); done = 0; past_bibliography = False
            annex_counter = 0; italic_count = 0; link_count = 0

            for kind, obj, in_cover in items:
                done += 1; pct = 5 + int(done / max(total, 1) * 60)
                if kind == 'para':
                    para = obj; is_bib = _is_biblio_title_para(para)
                    is_annex = _get_para_style_id(para) in _ANNEX_STYLE_IDS
                    hl_cnt = 1 if _has_hyperlinks(para) else 0

                    if in_cover and _all_runs_italic(para):
                        _notify(progress_callback, pct, "[Cover-italic] skip")
                    elif is_bib:
                        italic_count += len(_translate_para(para, tr))
                        past_bibliography = True
                    elif is_annex and not past_bibliography:
                        _translate_para(para, tr)
                        _fix_annex_style_para(para, chr(ord('A') + annex_counter))
                        annex_counter += 1
                    elif not _skip_paragraph(para, past_bibliography):
                        italic_count += len(_translate_para(para, tr))
                        link_count += hl_cnt
                elif kind == 'table':
                    if not past_bibliography: _translate_table(obj, tr)

            _notify(progress_callback, 66, "Em-dash bullets...")
            _convert_emdash_to_bullets(doc)
            _notify(progress_callback, 67, "Fix Note/Catatan...")
            _fix_all_notes(doc)

            if translate_headers:
                _notify(progress_callback, 70, "Translating headers/footers...")
                for s in doc.sections:
                    for hf in [s.header, s.footer, s.even_page_header, s.even_page_footer, s.first_page_header, s.first_page_footer]:
                        _translate_hf(hf, tr)

            _notify(progress_callback, 75, "Insert original...")
            _insert_original_before_bib(doc, orig_body_els, lambda p, m: _notify(progress_callback, 75 + int(p * 0.20), m))

            _notify(progress_callback, 96, "Sinkronisasi judul...")
            fid, fen = _extract_cover_titles(doc)
            if fid:
                _sync_foreword_title(doc, fid, fen)
                _sync_body_title(doc, fid)

            _notify(progress_callback, 97, "Saving...")
            doc.save(output_docx)
            
            summary = f"✅ Done!"
            if italic_count > 0: summary += f" Miring: {italic_count}."
            if link_count > 0: summary += f" Link: {link_count}."
            _notify(progress_callback, 100, summary)
            return True, output_docx

        except ImportError as e: return False, f"Dependensi: {e}"
        except Exception as e: return False, f"Engine9 Error: {str(e)}\n{traceback.format_exc()}"
