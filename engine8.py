"""
Engine8: DocxFinalTranslatorEngine
====================================
Menerjemahkan dokumen DOCX hasil pembangunan akhir dari bahasa asing
ke Bahasa Indonesia, kemudian menyisipkan konten asli (bahasa asing)
SEBELUM bagian Bibliografi.

Fitur utama:
  - Isi bibliografi TIDAK diterjemahkan.
  - Nomor pasal heading di bagian teks asli di-RESET ke 1.
  - Judul asing italic di cover TIDAK diterjemahkan.
  - **FITUR: Konversi em dash (—) menjadi Bullet Numbering Word (Flush Left).**
  - Heading ANNEX: rekonstruksi penuh XML.
    Format:
      Line 1: Annex A (Autonumber)
      Line 2: (informatif) -> Arial 11
      Line 3: [KOSONG] -> Jarak 1 Enter (Double Break)
      Line 4: Judul      -> Arial 11 Bold
  - **FITUR: Sisipkan Halaman "Introduction" sebelum konten asli.**
    Format: Arial, Center, Size 11, Bold.

Dependensi:
    pip install deep-translator python-docx lxml
"""

import re
import copy
import time
import traceback

from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree


# ─────────────────────────────────────────────────────────────────────────────
# NAMESPACE
# ─────────────────────────────────────────────────────────────────────────────

_NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W    = f'{{{_NS_W}}}'


# ─────────────────────────────────────────────────────────────────────────────
# KONSTANTA
# ─────────────────────────────────────────────────────────────────────────────

_RE_PURE_NUMBER = re.compile(
    r'^[\d\s\.\,\:\;\-\(\)\[\]\/\\\+\=\*\%\&\^\$\#\@\!\"\'`~<>{}|_]+$'
)
_RE_COPYRIGHT = re.compile(r'©|BSN\s*\d{4}', re.IGNORECASE)

_SKIP_STYLES = {
    'caption', 'header', 'footer',
    'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5',
    'table of figures', 'footnote text', 'endnote text', 'macro text',
}
_BIBLIO_TITLE_STYLES = {'biblioti', 'bibliotitle', 'bibliography title'}
_BIBLIO_KEYWORDS_EXACT = {
    'bibliografi', 'bibliography',
    'daftar acuan', 'daftar pustaka', 'daftar referensi',
}
_ANNEX_STYLE_IDS = {'ANNEX', 'Annex', 'annex'}
_HEADING_STYLES_WITH_NUM = {
    'Heading1', 'Heading2', 'Heading3',
    'ANNEX', 'a2', 'a3',
    'Heading4', 'Heading5', 'Heading6',
}
_TRANSLATE_DELAY = 0.15
_EM_DASH = '—'


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: FILTER
# ─────────────────────────────────────────────────────────────────────────────

def _skip_text(text: str) -> bool:
    t = text.strip()
    if len(t) < 3:                   return True
    if _RE_PURE_NUMBER.fullmatch(t): return True
    if _RE_COPYRIGHT.search(t):      return True
    return False

def _skip_paragraph(para, past_bibliography: bool = False) -> bool:
    if past_bibliography:
        return True
    if not para.text.strip():
        return True
    for tag in [f'{_W}drawing', f'{_W}pict']:
        if para._element.find('.//' + tag) is not None:
            return True
    style_name = (para.style.name or '').lower()
    if any(style_name.startswith(s) for s in _SKIP_STYLES):
        return True
    return False

def _is_biblio_title_para(para) -> bool:
    style_id = ''
    try:
        pStyle = para._element.find(f'{_W}pPr/{_W}pStyle')
        if pStyle is not None:
            style_id = pStyle.get(f'{_W}val', '').lower()
    except Exception:
        pass
    if style_id in _BIBLIO_TITLE_STYLES:
        return True
    txt = para.text.strip().lower()
    return bool(txt) and not txt[0].isdigit() and txt in _BIBLIO_KEYWORDS_EXACT

def _get_para_style_id(para) -> str:
    pStyle = para._element.find(f'{_W}pPr/{_W}pStyle')
    if pStyle is not None:
        return pStyle.get(f'{_W}val', '')
    return ''

def _get_style_id(el) -> str:
    pStyle = el.find(f'{_W}pPr/{_W}pStyle')
    return pStyle.get(f'{_W}val', '') if pStyle is not None else 'Normal'


# ─────────────────────────────────────────────────────────────────────────────
# FITUR 1: REKONSTRUKSI ANNEX (Dengan Jarak "Enter" yang Jelas)
# ─────────────────────────────────────────────────────────────────────────────

def _fix_annex_style_para(para) -> None:
    """
    Format Annex:
      Line 1: Annex A (Autonumber)
      Line 2: (informatif)
      Line 3: [KOSONG] -> Jarak 1 Enter
      Line 4: Judul
    """
    sid = _get_para_style_id(para)
    if sid not in _ANNEX_STYLE_IDS:
        return

    full_text = para.text.strip()
    if not full_text:
        return

    # 1. Cari dan Pisahkan Tag (informatif/normatif)
    tag_norm = None
    title_part = full_text
    
    # Daftar variasi tag yang mungkin
    tags_to_find = [
        '(informatif)', '(normatif)', 
        '(informative)', '(normative)',
        '(informasi)'
    ]
    
    for t in tags_to_find:
        # Cari posisi tag (case insensitive search)
        idx = full_text.lower().find(t)
        if idx != -1:
            # Tentukan tag standar output
            if 'norm' in t.lower():
                tag_norm = '(normatif)'
            else:
                tag_norm = '(informatif)'
            
            # Potong string: ambil bagian sebelum dan sesudah tag
            part_before = full_text[:idx]
            part_after = full_text[idx + len(t):]
            
            # Gabungkan sisa teks (jika ada teks sebelum/sesudah)
            title_part = part_before + " " + part_after
            break # Hanya proses tag pertama yang ditemukan

    # 2. Bersihkan Judul dari Prefix "Annex A"
    # Regex: Hapus kata "Annex" diikuti spasi dan huruf/angka/titik
    title_part = re.sub(r'^Annex\s+[A-Z0-9\.]+\s*', '', title_part, flags=re.IGNORECASE).strip()

    # 3. Hapus semua run lama
    pPr = para._element.find(f'{_W}pPr')
    for child in list(para._element):
        if child is not pPr:
            para._element.remove(child)

    # Helper XML
    def make_arial_run(text, is_bold=False):
        esc = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        b_tag = '<w:b/><w:bCs/>' if is_bold else ''
        return etree.fromstring(
            f'<w:r xmlns:w="{_NS_W}">'
            f'<w:rPr>'
            f'<w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
            f'{b_tag}'
            f'<w:sz w:val="22"/><w:szCs w:val="22"/>'
            f'</w:rPr>'
            f'<w:t xml:space="preserve">{esc}</w:t>'
            f'</w:r>'
        )

    def make_br_run():
        # Break standar
        return etree.fromstring(f'<w:r xmlns:w="{_NS_W}"><w:br/></w:r>')

    # 4. Susun Ulang dengan Jarak
    new_runs = []
    
    # A. Turun dari Autonumber (Break pertama)
    new_runs.append(make_br_run())
    
    # B. Tulis Tag (informatif) -> Arial 11
    if tag_norm:
        new_runs.append(make_arial_run(tag_norm, is_bold=False))
    
    # C. Tulis Judul
    if title_part:
        # Jarak 1 Enter = 2 Break. 
        # Satu break untuk turun ke baris baru dari tag.
        # Satu break lagi untuk membuat baris kosong (enter).
        new_runs.append(make_br_run()) # Turun baris
        new_runs.append(make_br_run()) # Buat spasi kosong
        
        # Tulis Judul
        new_runs.append(make_arial_run(title_part, is_bold=True))

    for run_el in new_runs:
        para._element.append(run_el)


# ─────────────────────────────────────────────────────────────────────────────
# FITUR 2: EM DASH TO BULLETS (Force Flush Left)
# ─────────────────────────────────────────────────────────────────────────────

def _get_or_create_emdash_numid(doc: Document) -> str:
    try:
        np = doc.part.numbering_part
        if np is None: return None
        nxml = np._element
        
        target_abstract_id = None
        for ab in nxml.findall(f'{_W}abstractNum'):
            for lvl in ab.findall(f'{_W}lvl'):
                txt_el = lvl.find(f'{_W}lvlText')
                if txt_el is not None and txt_el.get(f'{_W}val') == _EM_DASH:
                    target_abstract_id = ab.get(f'{_W}abstractNumId')
                    break
            if target_abstract_id: break
        
        if target_abstract_id:
            existing_nums = nxml.findall(f'{_W}num')
            max_num_id = max((int(n.get(f'{_W}numId', 0)) for n in existing_nums), default=0)
            new_num_id = str(max_num_id + 1)
            new_num = etree.fromstring(
                f'<w:num xmlns:w="{_NS_W}" w:numId="{new_num_id}">'
                f'<w:abstractNumId w:val="{target_abstract_id}"/>'
                f'</w:num>'
            )
            nxml.append(new_num)
            return new_num_id

        existing_abstracts = nxml.findall(f'{_W}abstractNum')
        max_abstract_id = max((int(a.get(f'{_W}abstractNumId', 0)) for a in existing_abstracts), default=0)
        new_abstract_id = str(max_abstract_id + 1)
        
        existing_nums = nxml.findall(f'{_W}num')
        max_num_id = max((int(n.get(f'{_W}numId', 0)) for n in existing_nums), default=0)
        new_num_id = str(max_num_id + 1)

        abstract_xml = f'''
        <w:abstractNum xmlns:w="{_NS_W}" w:abstractNumId="{new_abstract_id}">
            <w:multiLevelType w:val="hybridMultilevel"/>
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="bullet"/>
                <w:lvlText w:val="{_EM_DASH}"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="360" w:hanging="360"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
                </w:rPr>
            </w:lvl>
        </w:abstractNum>
        '''
        nxml.insert(0, etree.fromstring(abstract_xml))
        
        num_xml = f'''
        <w:num xmlns:w="{_NS_W}" w:numId="{new_num_id}">
            <w:abstractNumId w:val="{new_abstract_id}"/>
        </w:num>
        '''
        nxml.append(etree.fromstring(num_xml))
        
        return new_num_id

    except Exception as e:
        print(f"[Engine8] Error numbering: {e}")
        return None


def _convert_emdash_to_bullets(doc: Document) -> None:
    num_id = _get_or_create_emdash_numid(doc)
    if not num_id: return

    for para in doc.paragraphs:
        sid = _get_para_style_id(para)
        if sid in _ANNEX_STYLE_IDS or sid.startswith('Heading'):
            continue
            
        text = para.text.strip()
        if text.startswith(_EM_DASH):
            for r in para.runs:
                if _EM_DASH in r.text:
                    r.text = r.text.replace(_EM_DASH, "", 1).lstrip()
                    break
            
            pPr = para._element.find(f'{_W}pPr')
            if pPr is None:
                pPr = etree.SubElement(para._element, f'{_W}pPr')
            
            old_num = pPr.find(f'{_W}numPr')
            if old_num is not None: pPr.remove(old_num)
            
            new_num_pr = etree.fromstring(
                f'<w:numPr xmlns:w="{_NS_W}">'
                f'<w:ilvl w:val="0"/>'
                f'<w:numId w:val="{num_id}"/>'
                f'</w:numPr>'
            )
            pPr.insert(0, new_num_pr)

            old_ind = pPr.find(f'{_W}ind')
            if old_ind is not None: pPr.remove(old_ind)
            
            new_ind = etree.fromstring(
                f'<w:ind xmlns:w="{_NS_W}" w:left="360" w:hanging="360"/>'
            )
            pPr.insert(1, new_ind)


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: DETEKSI ITALIC & SECTION BREAK
# ─────────────────────────────────────────────────────────────────────────────

def _has_inline_sectpr(para) -> bool:
    pPr = para._element.find(f'{_W}pPr')
    if pPr is None: return False
    return pPr.find(f'{_W}sectPr') is not None

def _all_runs_italic(para) -> bool:
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs: return False
    para_style_italic = False
    try:
        if para.style and para.style.font and para.style.font.italic:
            para_style_italic = True
    except Exception: pass
    for run in text_runs:
        italic = False
        if run.font.italic is True: italic = True
        if not italic:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{_W}i')
                if i_el is not None:
                    val = i_el.get(f'{_W}val', 'true')
                    if val.lower() not in ('false', '0'): italic = True
        if not italic and para_style_italic:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{_W}i')
                if i_el is not None:
                    val = i_el.get(f'{_W}val', 'true')
                    if val.lower() not in ('false', '0'): italic = True
                else: italic = True
            else: italic = True
        if not italic: return False
    return True


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: BIBLIOGRAFI & AUTONUMBERING
# ─────────────────────────────────────────────────────────────────────────────

def _el_text(el) -> str:
    return ''.join(t.text or '' for t in el.findall(f'.//{_W}t')).strip()

def _is_bibliography_el(el) -> bool:
    if el.tag != f'{_W}p': return False
    sid = _get_style_id(el).lower()
    if sid in _BIBLIO_TITLE_STYLES: return True
    text = _el_text(el).lower().strip()
    if not text or text[0].isdigit(): return False
    return text in _BIBLIO_KEYWORDS_EXACT

def _find_bib_index(body_els: list) -> int:
    for i, el in enumerate(body_els):
        if _is_bibliography_el(el): return i
    return -1

def _find_content_start_before_bib(body_els: list, bib_idx: int) -> int:
    search_limit = bib_idx if bib_idx >= 0 else len(body_els)
    last_sectpr  = -1
    for i, el in enumerate(body_els):
        if i >= search_limit: break
        if el.tag != f'{_W}p': continue
        pPr = el.find(f'{_W}pPr')
        if pPr is not None and pPr.find(f'{_W}sectPr') is not None:
            last_sectpr = i
    return last_sectpr + 1

def _read_style_numpr(doc: Document) -> dict:
    result = {}
    try:
        styles_xml = doc.part.styles._element
        for style_el in styles_xml.findall(f'{_W}style'):
            sid = style_el.get(f'{_W}styleId', '')
            if sid not in _HEADING_STYLES_WITH_NUM: continue
            pPr = style_el.find(f'{_W}pPr')
            if pPr is None: continue
            numPr = pPr.find(f'{_W}numPr')
            if numPr is None: continue
            numId_el = numPr.find(f'{_W}numId')
            ilvl_el  = numPr.find(f'{_W}ilvl')
            nid  = numId_el.get(f'{_W}val', '0') if numId_el is not None else '0'
            ilvl = ilvl_el.get(f'{_W}val',  '0') if ilvl_el  is not None else '0'
            result[sid] = {'numId': nid, 'ilvl': ilvl}
    except Exception: pass
    return result

def _create_restart_numids(doc: Document, style_numpr: dict) -> dict:
    remapping = {}
    if not style_numpr: return remapping
    try:
        np   = doc.part.numbering_part
        nxml = np._element
        unique_numids = {info['numId'] for info in style_numpr.values()}
        existing = nxml.findall(f'{_W}num')
        max_id   = max((int(n.get(f'{_W}numId', 0)) for n in existing), default=0)
        for old_nid in unique_numids:
            if old_nid == '0': continue
            old_num_el = None
            for n in existing:
                if n.get(f'{_W}numId') == old_nid:
                    old_num_el = n
                    break
            if old_num_el is None: continue
            max_id += 1
            new_nid = str(max_id)
            new_num_el = copy.deepcopy(old_num_el)
            new_num_el.set(f'{_W}numId', new_nid)
            for lo in new_num_el.findall(f'{_W}lvlOverride'):
                new_num_el.remove(lo)
            override_el = etree.fromstring(
                f'<w:lvlOverride xmlns:w="{_NS_W}" w:ilvl="0">'
                f'<w:startOverride w:val="1"/>'
                f'</w:lvlOverride>'
            )
            new_num_el.append(override_el)
            nxml.append(new_num_el)
            remapping[old_nid] = new_nid
    except Exception: pass
    return remapping

def _apply_numpr_restart_to_headings(elements: list, style_numpr: dict, remapping: dict) -> None:
    if not remapping: return
    for el in elements:
        if el.tag != f'{_W}p': continue
        sid = _get_style_id(el)
        if sid not in style_numpr: continue
        info    = style_numpr[sid]
        old_nid = info['numId']
        ilvl    = info['ilvl']
        new_nid = remapping.get(old_nid)
        if not new_nid: continue
        pPr = el.find(f'{_W}pPr')
        if pPr is None:
            pPr = etree.SubElement(el, f'{_W}pPr')
            el.insert(0, pPr)
        old_numPr = pPr.find(f'{_W}numPr')
        if old_numPr is not None: pPr.remove(old_numPr)
        new_numPr = etree.fromstring(
            f'<w:numPr xmlns:w="{_NS_W}">'
            f'<w:ilvl w:val="{ilvl}"/>'
            f'<w:numId w:val="{new_nid}"/>'
            f'</w:numPr>'
        )
        pPr.insert(0, new_numPr)


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: ELEMEN PEMISAH
# ─────────────────────────────────────────────────────────────────────────────

def _page_break_para() -> etree._Element:
    return etree.fromstring(
        f'<w:p xmlns:w="{_NS_W}">'
        f'<w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
        f'<w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )

def _separator_heading(text: str) -> etree._Element:
    # Heading lama dengan border (masih disimpan jika dibutuhkan)
    esc = (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
    return etree.fromstring(
        f'<w:p xmlns:w="{_NS_W}">'
        f'<w:pPr>'
        f'<w:jc w:val="center"/>'
        f'<w:spacing w:before="280" w:after="280"/>'
        f'<w:pBdr>'
        f'<w:top    w:val="single" w:sz="12" w:space="4" w:color="4472C4"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="4472C4"/>'
        f'</w:pBdr>'
        f'</w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
        f'<w:b/><w:color w:val="4472C4"/>'
        f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{esc}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )

def _intro_heading() -> etree._Element:
    # Heading "Introduction" sesuai permintaan: Arial, Center, Size 11, BOLD
    return etree.fromstring(
        f'<w:p xmlns:w="{_NS_W}">'
        f'<w:pPr>'
        f'<w:jc w:val="center"/>'
        f'<w:spacing w:before="0" w:after="0"/>'
        f'</w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
        f'<w:b/><w:bCs/>' # <--- ADDED BOLD
        f'<w:sz w:val="22"/><w:szCs w:val="22"/>'
        f'</w:rPr>'
        f'<w:t>Introduction</w:t>'
        f'</w:r>'
        f'</w:p>'
    )

def _empty_para() -> etree._Element:
    return etree.fromstring(
        f'<w:p xmlns:w="{_NS_W}">'
        f'<w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
        f'</w:p>'
    )


# ─────────────────────────────────────────────────────────────────────────────
# CORE: TRANSLATION
# ─────────────────────────────────────────────────────────────────────────────

def _translate_para(para, tr, past_bibliography: bool = False) -> None:
    if _skip_paragraph(para, past_bibliography): return

    text_runs = [(i, r) for i, r in enumerate(para.runs)
                 if r.text and r.text.strip()]
    if not text_runs: return

    combined = ''.join(r.text for _, r in text_runs)
    if _skip_text(combined): return

    translated = tr.translate_one(combined.strip())
    time.sleep(_TRANSLATE_DELAY)
    if not translated or translated == combined: return

    _, first_run = text_runs[0]
    first_run.text = translated
    for _, run in text_runs[1:]:
        run.text = ''

def _translate_table(table, tr) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _translate_para(para, tr)

def _translate_hf(hf_part, tr) -> None:
    if hf_part is None: return
    try:
        for para in hf_part.paragraphs:
            if _RE_COPYRIGHT.search(para.text or ''): continue
            _translate_para(para, tr)
        for table in hf_part.tables:
            _translate_table(table, tr)
    except Exception: pass


# ─────────────────────────────────────────────────────────────────────────────
# CORE: INSERT ORIGINAL
# ─────────────────────────────────────────────────────────────────────────────

def _insert_original_before_bib(
    translated_doc:    Document,
    orig_body_els:     list,
    progress_callback=None,
) -> None:
    bib_in_orig   = _find_bib_index(orig_body_els)
    content_start = _find_content_start_before_bib(orig_body_els, bib_in_orig)
    content_end   = bib_in_orig if bib_in_orig >= 0 else len(orig_body_els)

    if bib_in_orig < 0 and content_end > content_start:
        if orig_body_els[content_end - 1].tag == f'{_W}sectPr':
            content_end -= 1

    els_to_insert = orig_body_els[content_start:content_end]

    if not els_to_insert:
        _notify(progress_callback, 50, "⚠️ Tidak ada konten asli yang bisa disisipkan.")
        return

    _notify(progress_callback, 10, f"Ditemukan {len(els_to_insert)} elemen.")

    style_numpr = _read_style_numpr(translated_doc)
    _notify(progress_callback, 20, f"Styles: {list(style_numpr.keys())}")

    remapping = _create_restart_numids(translated_doc, style_numpr)
    _notify(progress_callback, 30, f"Remapping: {remapping}")

    new_content_els = [copy.deepcopy(el) for el in els_to_insert]
    _apply_numpr_restart_to_headings(new_content_els, style_numpr, remapping)
    _notify(progress_callback, 40, "Reset heading numbers.")

    trans_body     = translated_doc.element.body
    trans_children = list(trans_body)
    bib_idx_trans  = _find_bib_index(trans_children)

    if bib_idx_trans < 0:
        bib_idx_trans = len(trans_children)
        for i in range(len(trans_children) - 1, -1, -1):
            if trans_children[i].tag == f'{_W}sectPr':
                bib_idx_trans = i
                break
        _notify(progress_callback, 45, "Bib tidak ditemukan → akhir.")
    else:
        _notify(progress_callback, 45, f"Bib posisi [{bib_idx_trans}].")

    # Susun urutan:
    # 1. Page Break (Halaman kosong di atasnya)
    # 2. Introduction (Heading) -> BOLD
    # 3. Page Break (Memulai konten di halaman baru)
    # 4. Konten Asli
    
    new_els = (
        [_page_break_para(),
         _intro_heading(),
         _page_break_para()]
        + new_content_els
        + [_empty_para()]
    )

    total = len(new_els)
    for offset, el in enumerate(new_els):
        trans_body.insert(bib_idx_trans + offset, el)
        if progress_callback and offset % 20 == 0:
            pct = 45 + int(offset / max(total, 1) * 55)
            progress_callback(pct, f"Menyisipkan... ({offset}/{total})")

    _notify(progress_callback, 100, f"✅ Selesai ({len(els_to_insert)} item).")


def _notify(cb, pct: int, msg: str) -> None:
    if cb:
        try: cb(pct, msg)
        except Exception: pass


# ─────────────────────────────────────────────────────────────────────────────
# TRANSLATOR WRAPPER
# ─────────────────────────────────────────────────────────────────────────────

class _Translator:
    def __init__(self, source: str = 'auto', target: str = 'id'):
        try:
            from deep_translator import GoogleTranslator
            self._cls = GoogleTranslator
            self.source = source
            self.target = target
        except ImportError:
            raise ImportError("Jalankan: pip install deep-translator")

    def translate_one(self, text: str) -> str:
        t = text.strip()
        if not t or _skip_text(t): return text
        try:
            result = self._cls(source=self.source, target=self.target).translate(t)
            return result if result else text
        except Exception:
            time.sleep(0.8)
            try:
                result = self._cls(source=self.source, target=self.target).translate(t)
                return result if result else text
            except Exception: return text


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENGINE CLASS
# ─────────────────────────────────────────────────────────────────────────────

class DocxFinalTranslatorEngine:
    """
    Engine 8 v16:
      - Annex: Robust string splitting.
      - Double Break for spacing.
      - "Introduction" page before original text (Arial, Center, 11pt, BOLD).
    """

    def __init__(self, source_lang: str = 'auto', target_lang: str = 'id'):
        self.source_lang = source_lang
        self.target_lang = target_lang

    def translate(
        self,
        input_docx:        str,
        output_docx:       str,
        progress_callback=None,
        translate_headers: bool = False,
    ) -> tuple[bool, str]:
        try:
            _notify(progress_callback, 2, "Snapshot dokumen...")
            doc_orig      = Document(input_docx)
            orig_body_els = [copy.deepcopy(el) for el in doc_orig.element.body]
            del doc_orig

            _notify(progress_callback, 5, "Init translator...")
            tr  = _Translator(source=self.source_lang, target=self.target_lang)
            doc = Document(input_docx)

            body     = doc.element.body
            para_map = {p._element: p for p in doc.paragraphs}
            tbl_map  = {t._element: t for t in doc.tables}

            items    = []
            sec_brks = 0
            COVER_END = 1

            for child in body:
                if child in para_map:
                    in_cover = (sec_brks < COVER_END)
                    items.append(('para', para_map[child], in_cover))
                    if _has_inline_sectpr(para_map[child]):
                        sec_brks += 1
                elif child in tbl_map:
                    in_cover = (sec_brks < COVER_END)
                    items.append(('table', tbl_map[child], in_cover))

            total             = len(items)
            done              = 0
            past_bibliography = False

            for kind, obj, in_cover in items:
                done += 1
                pct = 5 + int(done / max(total, 1) * 60)

                if kind == 'para':
                    para = obj
                    is_bib_heading = _is_biblio_title_para(para)
                    is_annex       = _get_para_style_id(para) in _ANNEX_STYLE_IDS

                    if in_cover and _all_runs_italic(para):
                        _notify(progress_callback, pct, "[Cover-italic] skip")

                    elif is_bib_heading:
                        _translate_para(para, tr, past_bibliography=False)
                        past_bibliography = True
                        _notify(progress_callback, pct, "[Bib-heading]")

                    elif is_annex and not past_bibliography:
                        _translate_para(para, tr, past_bibliography=False)
                        _fix_annex_style_para(para)
                        _notify(progress_callback, pct, "[ANNEX] fixed")

                    elif not _skip_paragraph(para, past_bibliography):
                        _translate_para(para, tr, past_bibliography=False)

                elif kind == 'table':
                    if not past_bibliography:
                        _translate_table(obj, tr)

            _notify(progress_callback, 66, "Formatting em-dash bullets...")
            _convert_emdash_to_bullets(doc)

            if translate_headers:
                _notify(progress_callback, 70, "Translating headers...")
                for section in doc.sections:
                    for hf in [
                        section.header,            section.footer,
                        section.even_page_header,  section.even_page_footer,
                        section.first_page_header, section.first_page_footer,
                    ]:
                        _translate_hf(hf, tr)

            _notify(progress_callback, 75, "Inserting original content...")
            
            def _insert_cb(pct_inner, msg):
                _notify(progress_callback, 75 + int(pct_inner * 0.20), msg)

            _insert_original_before_bib(
                translated_doc=doc,
                orig_body_els=orig_body_els,
                progress_callback=_insert_cb,
            )

            _notify(progress_callback, 97, "Saving...")
            doc.save(output_docx)
            _notify(progress_callback, 100, "✅ Done!")
            return True, output_docx

        except ImportError as e:
            return False, f"Dependensi tidak ditemukan: {e}\nJalankan: pip install deep-translator"
        except Exception as e:
            return False, f"Engine8 Error: {str(e)}\n{traceback.format_exc()}"
