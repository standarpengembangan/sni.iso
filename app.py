import streamlit as st
import os
import re
import time
import glob
import atexit
import threading
from io import BytesIO

# ─────────────────────────────────────────────────────────────────────────────
# AUTO-CLEANUP — pembersihan file temporer otomatis
# ─────────────────────────────────────────────────────────────────────────────

# Pola file temporer yang dibuat oleh aplikasi
_TEMP_PATTERNS = ["temp_main_*", "opt_*", "cover_*", "di_*", "pp_*", "ip_*", "ID_*"]
# Hapus file lebih lama dari N menit
_MAX_AGE_MINUTES = 30

def _cleanup_temp_files(max_age_minutes: int = _MAX_AGE_MINUTES, silent: bool = True):
    """Hapus semua file temporer yang lebih lama dari max_age_minutes."""
    deleted, freed = 0, 0
    cutoff = time.time() - (max_age_minutes * 60)
    for pattern in _TEMP_PATTERNS:
        for fpath in glob.glob(pattern):
            try:
                if os.path.isfile(fpath) and os.path.getmtime(fpath) < cutoff:
                    size = os.path.getsize(fpath)
                    os.remove(fpath)
                    deleted += 1
                    freed += size
            except Exception:
                pass
    if not silent and deleted > 0:
        mb = freed / (1024 * 1024)
        print(f"[AutoCleanup] {deleted} file dihapus, {mb:.2f} MB dibebaskan.")
    return deleted, freed

def _cleanup_session_files(session_state):
    """Hapus file milik sesi saat ini segera."""
    keys = ['_target_file', '_final_opt_file', '_final_tr_file']
    for k in keys:
        fpath = session_state.get(k)
        if fpath and os.path.isfile(fpath):
            try:
                os.remove(fpath)
            except Exception:
                pass

def _start_background_cleanup():
    """Jalankan cleanup berkala di background thread (tiap 15 menit)."""
    def _loop():
        while True:
            time.sleep(15 * 60)
            _cleanup_temp_files(silent=False)
    t = threading.Thread(target=_loop, daemon=True)
    t.start()

# Jalankan background cleanup sekali saat modul pertama kali diload
if 'bg_cleanup_started' not in st.session_state:
    _cleanup_temp_files(silent=True)   # bersihkan sisa sesi sebelumnya
    _start_background_cleanup()
    st.session_state['bg_cleanup_started'] = True

# Daftarkan cleanup saat proses Python berhenti (atexit)
atexit.register(_cleanup_temp_files, max_age_minutes=0, silent=False)

# --- IMPORT ENGINE ---
from engine2 import DocxOptimizerEngine
from engine4 import CoverPageEngine
from engine5 import DaftarIsiEngine
from engine6 import PrakataPendahuluanEngine
from engine7 import InfoPendukungEngine
from engine9 import CustomDictionary, DocxFinalTranslatorEngine

# --- KONFIGURASI HALAMAN ---
st.set_page_config(
    page_title="ISO Doc Master",
    page_icon="📑",
    layout="centered", # Centered untuk fokus tampilan
    initial_sidebar_state="collapsed" # Sidebar sembunyi default untuk clean look
)

# --- CSS CUSTOM ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap');

/* ══════════════════════════════════════════
   BASE
══════════════════════════════════════════ */
html, body, [class*="css"] {
    font-family: 'Outfit', sans-serif !important;
}
#MainMenu, footer, header { visibility: hidden; }

.stApp {
    background: #0d0f1a;
    background-image:
        radial-gradient(ellipse 80% 50% at 20% 10%, rgba(99,102,241,0.12) 0%, transparent 60%),
        radial-gradient(ellipse 60% 40% at 80% 80%, rgba(16,185,129,0.08) 0%, transparent 55%),
        radial-gradient(ellipse 50% 60% at 50% 50%, rgba(244,63,94,0.04) 0%, transparent 70%);
    min-height: 100vh;
}

/* ── Padding konten ── */
.block-container {
    padding-top: 2rem !important;
    padding-bottom: 3rem !important;
    max-width: 780px !important;
}

/* ══════════════════════════════════════════
   HEADER HERO
══════════════════════════════════════════ */
.app-header {
    position: relative;
    padding: 3rem 2rem 2.5rem;
    border-radius: 24px;
    text-align: center;
    margin-bottom: 1.8rem;
    overflow: hidden;
    background: linear-gradient(135deg, #1e1b4b 0%, #1e293b 50%, #0f172a 100%);
    border: 1px solid rgba(99,102,241,0.3);
    box-shadow: 0 0 60px rgba(99,102,241,0.15), 0 20px 40px rgba(0,0,0,0.4);
}
.app-header::before {
    content: '';
    position: absolute; inset: 0;
    background:
        radial-gradient(ellipse 60% 60% at 15% 40%, rgba(99,102,241,0.25) 0%, transparent 55%),
        radial-gradient(ellipse 40% 40% at 85% 20%, rgba(16,185,129,0.15) 0%, transparent 50%),
        radial-gradient(ellipse 30% 30% at 50% 90%, rgba(244,63,94,0.1) 0%, transparent 50%);
    pointer-events: none;
}
.app-header::after {
    content: '';
    position: absolute;
    top: -50%; left: -50%;
    width: 200%; height: 200%;
    background: repeating-linear-gradient(
        45deg,
        transparent,
        transparent 60px,
        rgba(255,255,255,0.012) 60px,
        rgba(255,255,255,0.012) 61px
    );
    pointer-events: none;
}
.app-header .badge {
    display: inline-block;
    background: rgba(99,102,241,0.2);
    border: 1px solid rgba(99,102,241,0.5);
    color: #a5b4fc;
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 2px;
    text-transform: uppercase;
    padding: 0.3rem 0.9rem;
    border-radius: 99px;
    margin-bottom: 1rem;
    position: relative;
}
.app-header h1 {
    margin: 0 0 0.5rem;
    font-size: 2.6rem;
    font-weight: 800;
    letter-spacing: -1px;
    line-height: 1.1;
    position: relative;
    background: linear-gradient(135deg, #e2e8f0 0%, #a5b4fc 50%, #6ee7b7 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
}
.app-header p {
    margin: 0;
    color: rgba(255,255,255,0.45);
    font-size: 0.9rem;
    font-weight: 400;
    position: relative;
}
.app-header .stats-row {
    display: flex;
    justify-content: center;
    gap: 2rem;
    margin-top: 1.5rem;
    position: relative;
}
.app-header .stat-item {
    text-align: center;
}
.app-header .stat-num {
    font-size: 1.3rem;
    font-weight: 700;
    color: #a5b4fc;
    line-height: 1;
}
.app-header .stat-lbl {
    font-size: 0.68rem;
    color: rgba(255,255,255,0.35);
    text-transform: uppercase;
    letter-spacing: 1px;
    margin-top: 0.2rem;
}
.app-header .stat-divider {
    width: 1px;
    background: rgba(255,255,255,0.1);
    align-self: stretch;
}

/* ══════════════════════════════════════════
   STATUS KAMUS — ganti st.success/warning
══════════════════════════════════════════ */
div[data-testid="stAlert"] {
    border-radius: 14px !important;
    border: none !important;
    font-size: 0.88rem !important;
    font-weight: 500 !important;
}
div[data-testid="stAlert"][data-baseweb="notification"] {
    background: rgba(16,185,129,0.12) !important;
    border-left: 3px solid #10b981 !important;
    color: #6ee7b7 !important;
}

/* ══════════════════════════════════════════
   SECTION LABEL
══════════════════════════════════════════ */
.section-label {
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 2px;
    text-transform: uppercase;
    color: rgba(165,180,252,0.6);
    margin: 1.6rem 0 0.8rem;
}

/* ══════════════════════════════════════════
   FILE UPLOADER
══════════════════════════════════════════ */
section[data-testid="stFileUploaderDropzone"] {
    background: rgba(30, 27, 75, 0.4) !important;
    border: 2px dashed rgba(99,102,241,0.35) !important;
    border-radius: 16px !important;
    transition: all 0.25s ease;
    backdrop-filter: blur(10px);
}
section[data-testid="stFileUploaderDropzone"]:hover {
    border-color: rgba(99,102,241,0.7) !important;
    background: rgba(99,102,241,0.08) !important;
    box-shadow: 0 0 20px rgba(99,102,241,0.15);
}
section[data-testid="stFileUploaderDropzone"] p,
section[data-testid="stFileUploaderDropzone"] span {
    color: rgba(255,255,255,0.5) !important;
}
div[data-testid="stFileUploaderFile"] {
    background: rgba(99,102,241,0.1) !important;
    border: 1px solid rgba(99,102,241,0.3) !important;
    border-radius: 10px !important;
    color: #c7d2fe !important;
}

/* ══════════════════════════════════════════
   INPUT FIELDS
══════════════════════════════════════════ */
.stTextInput label, .stSelectbox label {
    color: rgba(255,255,255,0.55) !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.3px;
}
.stTextInput input {
    background: rgba(15,23,42,0.8) !important;
    border: 1.5px solid rgba(99,102,241,0.25) !important;
    border-radius: 12px !important;
    color: #e2e8f0 !important;
    font-family: 'Outfit', sans-serif !important;
    font-size: 0.9rem !important;
    transition: border-color 0.2s, box-shadow 0.2s;
}
.stTextInput input:focus {
    border-color: rgba(99,102,241,0.7) !important;
    box-shadow: 0 0 0 3px rgba(99,102,241,0.15) !important;
}
.stTextInput input::placeholder { color: rgba(255,255,255,0.2) !important; }

/* Selectbox */
.stSelectbox > div > div {
    background: rgba(15,23,42,0.8) !important;
    border: 1.5px solid rgba(99,102,241,0.25) !important;
    border-radius: 12px !important;
    color: #e2e8f0 !important;
}
.stSelectbox svg { color: rgba(165,180,252,0.6) !important; }

/* ══════════════════════════════════════════
   TOMBOL PROSES
══════════════════════════════════════════ */
.stButton > button {
    height: 3.2rem !important;
    border-radius: 14px !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
    letter-spacing: 0.3px !important;
    border: none !important;
    background: linear-gradient(135deg, #6366f1 0%, #4f46e5 50%, #4338ca 100%) !important;
    color: white !important;
    box-shadow: 0 4px 20px rgba(99,102,241,0.4), 0 1px 0 rgba(255,255,255,0.1) inset !important;
    transition: all 0.2s ease !important;
    position: relative;
    overflow: hidden;
}
.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 28px rgba(99,102,241,0.5), 0 1px 0 rgba(255,255,255,0.1) inset !important;
    background: linear-gradient(135deg, #818cf8 0%, #6366f1 50%, #4f46e5 100%) !important;
}
.stButton > button:active {
    transform: translateY(0) !important;
    box-shadow: 0 2px 10px rgba(99,102,241,0.3) !important;
}

/* ══════════════════════════════════════════
   PROGRESS
══════════════════════════════════════════ */
.stProgress > div {
    background: rgba(255,255,255,0.07) !important;
    border-radius: 99px !important;
    height: 8px !important;
}
.stProgress > div > div {
    background: linear-gradient(90deg, #6366f1, #818cf8, #10b981) !important;
    border-radius: 99px !important;
    box-shadow: 0 0 10px rgba(99,102,241,0.5);
    transition: width 0.4s ease !important;
}
div[data-testid="stProgressText"] {
    color: rgba(165,180,252,0.8) !important;
    font-size: 0.82rem !important;
}

/* ══════════════════════════════════════════
   TIMER
══════════════════════════════════════════ */
.timer-text {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.82rem;
    color: rgba(110,231,183,0.7);
    text-align: center;
    letter-spacing: 1px;
    margin: 0.4rem 0;
}

/* ══════════════════════════════════════════
   SPINNER
══════════════════════════════════════════ */
div[data-testid="stSpinner"] > div {
    color: #a5b4fc !important;
}

/* ══════════════════════════════════════════
   RESULT CARD
══════════════════════════════════════════ */
.result-panel {
    background: linear-gradient(135deg, rgba(30,27,75,0.6) 0%, rgba(15,23,42,0.8) 100%);
    border: 1px solid rgba(99,102,241,0.25);
    border-radius: 20px;
    padding: 2rem;
    text-align: center;
    backdrop-filter: blur(12px);
    box-shadow: 0 20px 40px rgba(0,0,0,0.3), 0 0 0 1px rgba(255,255,255,0.04) inset;
    margin: 1rem 0;
}
.result-panel .check-icon {
    font-size: 2.5rem;
    display: block;
    margin-bottom: 0.5rem;
    filter: drop-shadow(0 0 10px rgba(16,185,129,0.6));
}
.result-panel h3 {
    color: #e2e8f0;
    font-size: 1.2rem;
    font-weight: 700;
    margin: 0 0 0.3rem;
}
.result-panel .time-badge {
    display: inline-block;
    background: rgba(16,185,129,0.15);
    border: 1px solid rgba(16,185,129,0.3);
    color: #6ee7b7;
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.8rem;
    padding: 0.25rem 0.8rem;
    border-radius: 99px;
    margin-bottom: 1.5rem;
}

/* ══════════════════════════════════════════
   DOWNLOAD BUTTONS
══════════════════════════════════════════ */
.stDownloadButton > button {
    border-radius: 12px !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    height: 3rem !important;
    transition: all 0.2s ease !important;
    border: 1.5px solid rgba(99,102,241,0.3) !important;
    background: rgba(30,27,75,0.6) !important;
    color: #c7d2fe !important;
    backdrop-filter: blur(8px);
    box-shadow: 0 2px 12px rgba(0,0,0,0.2) !important;
}
.stDownloadButton > button:hover {
    background: rgba(99,102,241,0.2) !important;
    border-color: rgba(99,102,241,0.6) !important;
    color: #e0e7ff !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(99,102,241,0.25) !important;
}

/* ══════════════════════════════════════════
   AI MENU EXPANDER — ungu kebiruan
══════════════════════════════════════════ */
div[data-testid="stExpander"] {
    border: 1.5px solid rgba(99,102,241,0.55) !important;
    border-radius: 14px !important;
    box-shadow: 0 4px 20px rgba(99,102,241,0.3), 0 1px 0 rgba(255,255,255,0.06) inset !important;
    overflow: hidden !important;
    background: rgba(20,18,50,0.6) !important;
}
div[data-testid="stExpander"]:hover {
    border-color: rgba(129,140,248,0.8) !important;
    box-shadow: 0 8px 28px rgba(99,102,241,0.45) !important;
}

/* Header / tombol expander */
div[data-testid="stExpander"] > details > summary,
div[data-testid="stExpander"] details summary {
    background: linear-gradient(135deg, #6366f1 0%, #4f46e5 55%, #4338ca 100%) !important;
    border-radius: 12px !important;
    padding: 0.85rem 1.2rem !important;
    cursor: pointer !important;
    box-shadow: 0 4px 15px rgba(99,102,241,0.4) !important;
    list-style: none !important;
}
div[data-testid="stExpander"] details summary:hover {
    background: linear-gradient(135deg, #818cf8 0%, #6366f1 55%, #4f46e5 100%) !important;
    box-shadow: 0 6px 22px rgba(99,102,241,0.55) !important;
}

/* Semua teks di dalam summary */
div[data-testid="stExpander"] details summary *,
div[data-testid="stExpander"] details summary p,
div[data-testid="stExpander"] details summary span,
div[data-testid="stExpander"] details summary div,
div[data-testid="stExpander"] details summary label {
    color: #ffffff !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
}

/* Ikon panah */
div[data-testid="stExpander"] details summary svg {
    color: #ffffff !important;
    stroke: #ffffff !important;
    fill: #ffffff !important;
}

/* Konten dalam */
div[data-testid="stExpander"] details > div,
div[data-testid="stExpander"] .streamlit-expanderContent {
    background: transparent !important;
    padding-top: 0.6rem !important;
}

/* ══════════════════════════════════════════
   DIVIDER
══════════════════════════════════════════ */
hr {
    border: none !important;
    border-top: 1px solid rgba(255,255,255,0.07) !important;
    margin: 1.5rem 0 !important;
}

/* ══════════════════════════════════════════
   FOOTER
══════════════════════════════════════════ */
.footer {
    text-align: center;
    color: rgba(255,255,255,0.2);
    font-size: 0.75rem;
    padding: 2rem 0 1rem;
    letter-spacing: 0.5px;
}
.footer span { color: rgba(99,102,241,0.5); }

/* ══════════════════════════════════════════
   STATUS PILL — di dalam header
══════════════════════════════════════════ */
.status-pill {
    display: inline-flex;
    align-items: center;
    gap: 0.45rem;
    font-size: 0.8rem;
    font-weight: 600;
    padding: 0.38rem 1.1rem;
    border-radius: 99px;
    margin-top: 1.2rem;
    position: relative;
    letter-spacing: 0.2px;
}
.status-ready {
    background: rgba(16,185,129,0.15);
    border: 1px solid rgba(16,185,129,0.45);
    color: #6ee7b7;
}
.status-warn {
    background: rgba(245,158,11,0.13);
    border: 1px solid rgba(245,158,11,0.4);
    color: #fcd34d;
}
.status-dot {
    width: 7px; height: 7px;
    border-radius: 50%;
    display: inline-block;
    flex-shrink: 0;
}
.status-ready .status-dot {
    background: #10b981;
    box-shadow: 0 0 6px rgba(16,185,129,0.8);
    animation: pulse-green 2s infinite;
}
.status-warn .status-dot {
    background: #f59e0b;
    box-shadow: 0 0 6px rgba(245,158,11,0.8);
}
@keyframes pulse-green {
    0%, 100% { opacity: 1; transform: scale(1); }
    50% { opacity: 0.5; transform: scale(1.35); }
}

/* ══════════════════════════════════════════
   GENERAL TEXT FIX
══════════════════════════════════════════ */
p, li, span, div { color: inherit; }
.stApp p, .stApp div, .stApp label { color: rgba(255,255,255,0.75); }
</style>
""", unsafe_allow_html=True)


# --- KONSTANTA STANDAR ISO ---
ISO_FONT_NAME = "Arial"
ISO_FONT_SIZE = 11
LANG_OPTIONS = {
    "auto": "🔍 Deteksi Otomatis", "en": "🇬🇧 Inggris",
    "fr": "🇫🇷 Prancis", "de": "🇩🇪 Jerman", "es": "🇪🇸 Spanyol",
    "it": "🇮🇹 Italia", "nl": "🇳🇱 Belanda", "pt": "🇵🇹 Portugis",
    "ru": "🇷🇺 Rusia", "ja": "🇯🇵 Jepang", "zh-CN": "🇨🇳 Mandarin",
    "ko": "🇰🇷 Korea", "ar": "🇸🇦 Arab",
}

# --- HELPER FUNGSI ---
def _parse_doc_structure(docx_path: str) -> list:
    """Parsing dokumen menjadi daftar section: [{heading, level, paragraphs}]"""
    try:
        from docx import Document as _Doc
        doc = _Doc(docx_path)
        sections, current = [], {"heading": "Pembukaan", "level": 0, "paragraphs": []}
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            style = para.style.name.lower()
            if style.startswith("heading"):
                if current["paragraphs"]:
                    sections.append(current)
                try:
                    lvl = int(style.replace("heading", "").strip())
                except Exception:
                    lvl = 1
                current = {"heading": txt, "level": lvl, "paragraphs": []}
            else:
                current["paragraphs"].append(txt)
        if current["paragraphs"] or current["heading"] != "Pembukaan":
            sections.append(current)
        return sections
    except Exception:
        return []

def get_elapsed_str(start_time: float) -> str:
    elapsed = time.time() - start_time
    if elapsed < 60:
        return f"{elapsed:.1f} detik"
    else:
        mins = int(elapsed // 60)
        secs = elapsed % 60
        return f"{mins} menit {secs:.1f} detik"

def extract_titles_from_docx(docx_path: str):
    try:
        from docx import Document
        doc = Document(docx_path)
        candidates = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) < 5: continue
            is_heading = para.style.name.lower().startswith('heading')
            max_size = 0
            is_italic = False
            for run in para.runs:
                if run.text.strip():
                    sz = run.font.size
                    if sz: max_size = max(max_size, sz.pt if hasattr(sz, 'pt') else sz / 12700)
                    if run.font.italic: is_italic = True
            if is_heading or max_size >= 12:
                candidates.append({'text': text, 'size': max_size})
            if len(candidates) >= 10: break
        if not candidates: return "", ""
        title_id = candidates[0]['text']
        title_en = title_id
        return title_id, title_en
    except Exception: return "", ""

# --- INISIALISASI ENGINE ---
@st.cache_resource
def load_engines():
    e2 = DocxOptimizerEngine()
    e4 = CoverPageEngine()
    e5 = DaftarIsiEngine()
    e6 = PrakataPendahuluanEngine()
    e7 = InfoPendukungEngine()
    return e2, e4, e5, e6, e7

engine2, engine4, engine5, engine6, engine7 = load_engines()

# --- AUTO-LOAD KAMUS ---
@st.cache_resource(show_spinner=False)
def _load_kamus_from_sheet():
    d = CustomDictionary()
    count = d.load_defaults()
    return d, count

if 'custom_dict' not in st.session_state:
    _d, _n = _load_kamus_from_sheet()
    st.session_state['custom_dict'] = _d
    st.session_state['kamus_count'] = _n

_kamus = st.session_state.get('custom_dict')
_count = len(_kamus) if _kamus else 0

# --- HALAMAN UTAMA ---

# --- HEADER ---
_status_html = (
    f"""<div class="status-pill status-ready">
        <span class="status-dot"></span>
        Sistem Siap &nbsp;·&nbsp; Kamus Aktif: <strong>{_count} istilah</strong>
    </div>"""
    if _count > 0 else
    """<div class="status-pill status-warn">
        <span class="status-dot"></span>
        Kamus Tidak Aktif
    </div>"""
)

st.markdown(f"""
    <div class="app-header">
        <div class="badge">Generator RSNI</div>
        <h1>📑 ISO to SNI Master Formating</h1>
        <p>Memformat & menerjemahan Dokumen Standar ISO menjadi Draft RSNI secara otomatis</p>
        <div class="stats-row">
            <div class="stat-item">
                <div class="stat-num">6</div>
                <div class="stat-lbl">Engine</div>
            </div>
            <div class="stat-divider"></div>
            <div class="stat-item">
                <div class="stat-num">{_count if _count > 0 else '—'}</div>
                <div class="stat-lbl">Istilah Kamus</div>
            </div>
            <div class="stat-divider"></div>
            <div class="stat-item">
                <div class="stat-num">13</div>
                <div class="stat-lbl">Bahasa</div>
            </div>
        </div>
        {_status_html}
    </div>
""", unsafe_allow_html=True)

import datetime
_tahun = str(datetime.date.today().year)

# --- FORM INPUT ---
st.markdown('<div class="section-label">📂 Upload Dokumen ISO</div>', unsafe_allow_html=True)
uploaded_file = st.file_uploader("Upload file .docx di sini atau klik Browse", type=["docx"], key="upl_main", label_visibility="collapsed")

st.markdown('<div class="section-label">⚙️ Pengaturan</div>', unsafe_allow_html=True)
col_set1, col_set2 = st.columns([3, 2])
with col_set1:
    doc_title = st.text_input("📄 Masukan No. SNI", value="SNI ISO xxxxx-x:xxxx", key="title_main")
with col_set2:
    src_lang = st.selectbox(
        "🌐 Bahasa Sumber",
        options=list(LANG_OPTIONS.keys()),
        index=0,
        format_func=lambda x: LANG_OPTIONS[x],
        key="lang_main"
    )

btn_process = st.button("🚀 Proses Formating & Terjemahan", key="btn_main", use_container_width=True)


# ─────────────────────────────────────────────────────────────────────────────
# LOGIC EXECUTION
# ─────────────────────────────────────────────────────────────────────────────

if btn_process:
    if uploaded_file:
        target_file = f"temp_main_{uploaded_file.name}"
        with open(target_file, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        st.session_state['_run_process'] = True
        st.session_state['_target_file'] = target_file
        st.session_state['_doc_title'] = doc_title
        st.session_state['_src_lang'] = src_lang
        st.rerun()
    else:
        st.warning("Silakan upload file terlebih dahulu.")


if st.session_state.get('_run_process') and st.session_state.get('_target_file'):
    target_file = st.session_state['_target_file']
    doc_title_val = st.session_state['_doc_title']
    src_lang_val = st.session_state['_src_lang']
    
    # UI Progress — 3 elemen terpisah agar tidak saling tumpuk
    status_placeholder = st.empty()
    progress_bar = st.progress(0)
    time_placeholder = st.empty()
    start_time = time.time()
    
    # Helper Update UI
    def update_ui(pct, msg):
        status_placeholder.markdown(
            f'<div style="font-size:0.85rem; color:rgba(165,180,252,0.85); '
            f'font-family:\'Outfit\',sans-serif; font-weight:500; margin-bottom:0.3rem;">'
            f'⚡ {msg}</div>',
            unsafe_allow_html=True
        )
        progress_bar.progress(pct)
        time_placeholder.markdown(
            f'<div class="timer-text">⏱ {get_elapsed_str(start_time)}</div>',
            unsafe_allow_html=True
        )

    # Pipeline Optimasi
    def run_optimization(input_file, doc_title):
        copyright_text = f"© BSN {_tahun}"
        cover_settings = {
            "sni_number": doc_title if doc_title else "SNI ISO XXXXX:20XX",
            "bsn_year": _tahun, "ics_number": "XX.XXX.XX", "ref_standard": "",
        }

        # 1. Engine 2
        update_ui(5, "[1/6] Format Dasar...")
        output_file = f"opt_{os.path.basename(input_file)}"
        success, msg = engine2.process(input_file, output_file, ISO_FONT_NAME, ISO_FONT_SIZE, enable_headers=True, doc_title=doc_title, copyright_text=copyright_text)
        if not success: raise Exception(f"Engine 2: {msg}")
        final_file = output_file
        auto_title_id, auto_title_en = extract_titles_from_docx(output_file)

        # 2. Engine 4
        update_ui(20, "[2/6] Cover...")
        cover_out = f"cover_{os.path.basename(final_file)}"
        if engine4.prepend_cover(input_docx=final_file, output_docx=cover_out, sni_number=cover_settings["sni_number"], bsn_year=cover_settings["bsn_year"], title_id=auto_title_id, title_en=auto_title_en, ref_standard=cover_settings["ref_standard"], ics_number=cover_settings["ics_number"])[0]:
            final_file = cover_out

        # 3. Engine 5
        update_ui(35, "[3/6] Daftar Isi...")
        di_out = f"di_{os.path.basename(final_file)}"
        if engine5.insert(input_docx=final_file, output_docx=di_out, doc_title=cover_settings["sni_number"], copyright_text=f"©BSN {cover_settings['bsn_year']}")[0]:
            final_file = di_out

        # 4. Engine 6
        update_ui(50, "[4/6] Prakata...")
        pp_out = f"pp_{os.path.basename(final_file)}"
        ref_std = re.sub(r'^SNI\s+', '', cover_settings["sni_number"]).strip()
        if engine6.insert(input_docx=final_file, output_docx=pp_out, sni_number=cover_settings["sni_number"], title_id=auto_title_id or 'Judul ID', title_en=auto_title_en or 'Title EN', ref_standard=ref_std, bsn_year=cover_settings["bsn_year"])[0]:
            final_file = pp_out

        # 5. Engine 7
        update_ui(65, "[5/6] Info Pendukung...")
        ip_out = f"ip_{os.path.basename(final_file)}"
        if engine7.append(input_docx=final_file, output_docx=ip_out)[0]:
            final_file = ip_out
        
        return final_file

    try:
        # --- LANGKAH 1: OPTIMASI ---
        final_opt_file = run_optimization(target_file, doc_title_val)
        st.session_state['_final_opt_file'] = final_opt_file
        
        # --- LANGKAH 2: TERJEMAHAN ---
        update_ui(75, "[6/6] Terjemahan...")
        tr_out = f"ID_{os.path.basename(final_opt_file)}"
        
        _engine9 = DocxFinalTranslatorEngine(source_lang=src_lang_val, target_lang='id', custom_dict=st.session_state.get('custom_dict'))
        
        def _cb_tr(pct, msg):
            final_pct = 75 + int(pct * 0.25)
            update_ui(final_pct, f"[Translate] {msg}")

        ok_tr, _ = _engine9.translate(input_docx=final_opt_file, output_docx=tr_out, progress_callback=_cb_tr, translate_headers=False)
        
        if ok_tr:
            update_ui(100, "✅ Selesai!")
            st.session_state['_final_tr_file'] = tr_out
            st.session_state['_final_time'] = get_elapsed_str(start_time)
            st.session_state['_show_results'] = True
            # Parse dokumen langsung agar chat langsung siap setelah rerun
            st.session_state['_doc_sections'] = _parse_doc_structure(tr_out)
        else:
            raise Exception("Terjemahan gagal.")

    except Exception as e:
        st.error(f"❌ Error Proses: {e}")
        st.session_state['_run_process'] = False
        st.session_state['_show_results'] = False
    finally:
        if st.session_state.get('_show_results'):
            st.session_state['_run_process'] = False
            st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# TAMPILKAN HASIL AKHIR
# ─────────────────────────────────────────────────────────────────────────────

if st.session_state.get('_show_results'):
    st.divider()
    final_time = st.session_state.get('_final_time', '-')
    st.markdown(
        f"""<div class="result-panel">
            <span class="check-icon">✅</span>
            <h3>Dokumen Berhasil Diproses</h3>
            <div class="time-badge">⏱ {final_time}</div>
        </div>""",
        unsafe_allow_html=True
    )

    col_res1, col_res2 = st.columns(2)

    opt_file = st.session_state.get('_final_opt_file')
    if opt_file and os.path.exists(opt_file):
        with col_res1:
            with open(opt_file, "rb") as f:
                st.download_button(
                    label="📄 Download Hasil Formating",
                    data=f,
                    file_name="ISO_Fixed_Document.docx",
                    use_container_width=True
                )

    tr_file = st.session_state.get('_final_tr_file')
    if tr_file and os.path.exists(tr_file):
        with col_res2:
            with open(tr_file, "rb") as f:
                st.download_button(
                    label="🌐 Download Terjemahan ID",
                    data=f,
                    file_name=f"ID_{os.path.basename(opt_file) if opt_file else 'document.docx'}",
                    use_container_width=True
                )

    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
    if st.button("🔄 Proses File Baru", key="reset", use_container_width=True):
        # Hapus file sesi ini segera sebelum reset
        _cleanup_session_files(st.session_state)
        for k in ['_show_results', '_final_opt_file', '_final_tr_file', '_final_time', '_run_process',
                  '_doc_text', '_chat_history', '_doc_sections', '_target_file']:
            if k in st.session_state: del st.session_state[k]
        st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# MESIN ANALISIS LOKAL — 100% offline, tidak ada data keluar
# ─────────────────────────────────────────────────────────────────────────────

def _local_answer(query: str, sections: list, history: list) -> str:
    """Mesin jawab lokal berbasis pencarian dan ekstraksi dari struktur dokumen."""
    import difflib

    q = query.lower().strip()
    words = re.findall(r'\w+', q)

    # ── 1. Deteksi intent ─────────────────────────────────────────────────────

    # Ringkasan
    is_summary = any(w in q for w in [
        'ringkas', 'rangkum', 'ringkasan', 'rangkuman', 'isi dokumen',
        'gambaran', 'overview', 'tentang apa', 'dokumen ini', 'keseluruhan'
    ])
    # Daftar heading
    is_list_sections = any(w in q for w in [
        'daftar bab', 'daftar bagian', 'bagian apa', 'bab apa', 'struktur',
        'apa saja bagian', 'apa saja bab', 'daftar isi', 'section'
    ])
    # Definisi / istilah
    is_definition = any(p in q for p in [
        'apa itu', 'apa yang dimaksud', 'definisi', 'pengertian', 'artinya',
        'maksudnya', 'jelaskan', 'explain'
    ])
    # Cari / di mana
    is_search = any(p in q for p in [
        'di mana', 'dimana', 'cari', 'temukan', 'ada di', 'letak',
        'sebutkan', 'mention', 'terdapat', 'berisi'
    ])
    # Referensi / standar
    is_ref = any(w in q for w in [
        'referensi', 'acuan', 'standar', 'normatif', 'bibliography',
        'pustaka', 'rujukan', 'iso', 'sni', 'iec'
    ])
    # Pasal / klausul tertentu
    clause_match = re.search(r'(?:bab|bagian|klausul|pasal|sub|annex|lampiran|section)\s*[\d\.]+', q)

    # ── 2. Bangun jawaban ─────────────────────────────────────────────────────

    def _section_text(s, max_chars=600):
        body = " ".join(s["paragraphs"])
        return body[:max_chars] + ("..." if len(body) > max_chars else "")

    def _score(s, kws):
        """Skor relevansi section berdasarkan kemunculan keyword."""
        txt = (s["heading"] + " " + " ".join(s["paragraphs"])).lower()
        return sum(2 if w in s["heading"].lower() else (1 if w in txt else 0) for w in kws)

    if not sections:
        return "⚠️ Dokumen tidak dapat dianalisis. Pastikan file .docx valid."

    # Ringkasan keseluruhan
    if is_summary:
        headings = [f"**{s['heading']}**" for s in sections if s['level'] <= 2][:12]
        total_para = sum(len(s['paragraphs']) for s in sections)
        intro = _section_text(sections[0], 400) if sections else ""
        return (
            f"📄 **Ringkasan Dokumen**\n\n"
            f"Dokumen ini terdiri dari **{len(sections)} bagian** dengan total **{total_para} paragraf**.\n\n"
            f"**Bagian utama:**\n" + "\n".join(f"- {h}" for h in headings) +
            (f"\n\n**Pembukaan:**\n{intro}" if intro else "")
        )

    # Daftar section/bab
    if is_list_sections:
        lines = []
        for s in sections:
            indent = "  " * max(0, s['level'] - 1)
            lines.append(f"{indent}{'#' * s['level']} {s['heading']}")
        return "📋 **Struktur Dokumen:**\n\n```\n" + "\n".join(lines[:40]) + "\n```"

    # Referensi & standar
    if is_ref:
        ref_secs = [s for s in sections if any(
            w in s['heading'].lower() for w in ['referensi', 'acuan', 'normatif', 'bibliography', 'pustaka', 'standar']
        )]
        if ref_secs:
            results = []
            for s in ref_secs:
                results.append(f"**{s['heading']}**\n{_section_text(s, 800)}")
            return "📚 **Referensi & Acuan Normatif:**\n\n" + "\n\n---\n\n".join(results)
        # Cari ISO/SNI/IEC dalam seluruh dokumen
        found = []
        for s in sections:
            for p in s['paragraphs']:
                if re.search(r'\b(ISO|SNI|IEC|IEEE)\s*[\d\-:]+', p):
                    found.append(f"- _{s['heading']}:_ {p[:200]}")
        if found:
            return "📚 **Referensi standar yang ditemukan:**\n\n" + "\n".join(found[:15])
        return "ℹ️ Tidak ditemukan bagian referensi atau acuan normatif dalam dokumen ini."

    # Klausul / bab spesifik
    if clause_match:
        target = clause_match.group(0).lower()
        kws = re.findall(r'\w+', target)
        scored = sorted(sections, key=lambda s: _score(s, kws), reverse=True)
        best = scored[:3]
        if best and _score(best[0], kws) > 0:
            results = []
            for s in best:
                if _score(s, kws) > 0:
                    results.append(f"**{s['heading']}**\n{_section_text(s, 700)}")
            return f"📖 **Hasil pencarian '{target}':**\n\n" + "\n\n---\n\n".join(results)

    # Definisi / jelaskan istilah
    if is_definition:
        # Ekstrak kata kunci utama dari pertanyaan (buang stop words)
        stop = {'apa','itu','yang','dimaksud','definisi','pengertian','artinya',
                'maksudnya','jelaskan','dengan','dari','dan','di','ke','adalah'}
        kws = [w for w in words if w not in stop and len(w) > 2]
        if kws:
            scored = sorted(sections, key=lambda s: _score(s, kws), reverse=True)
            best = [s for s in scored if _score(s, kws) > 0][:3]
            if best:
                results = []
                for s in best:
                    # Cari paragraf spesifik yang mengandung kata kunci
                    relevant_paras = [p for p in s['paragraphs']
                                      if any(k in p.lower() for k in kws)][:3]
                    body = " ".join(relevant_paras) if relevant_paras else _section_text(s, 500)
                    results.append(f"**{s['heading']}**\n{body[:600]}")
                return f"🔍 **Penjelasan '{' '.join(kws)}':**\n\n" + "\n\n---\n\n".join(results)

    # Pencarian umum / cari kata kunci
    stop_general = {'apa','ada','di','ke','dan','atau','yang','adalah','ini','itu',
                    'dengan','untuk','dari','pada','dalam','tidak','bisa','cara',
                    'bagaimana','berapa','siapa','kapan','dimana','mana','sebutkan',
                    'cari','temukan','terdapat','berisi','tentang','mengenai'}
    kws = [w for w in words if w not in stop_general and len(w) > 2]

    if kws:
        scored = [(s, _score(s, kws)) for s in sections]
        scored = sorted(scored, key=lambda x: x[1], reverse=True)
        best = [(s, sc) for s, sc in scored if sc > 0][:4]

        if best:
            results = []
            for s, sc in best:
                relevant_paras = [p for p in s['paragraphs']
                                  if any(k in p.lower() for k in kws)][:2]
                body = " ".join(relevant_paras) if relevant_paras else _section_text(s, 400)
                results.append(f"**{s['heading']}**\n{body[:500]}")
            return (
                f"🔍 **Hasil pencarian '{query}'** — ditemukan di {len(best)} bagian:\n\n" +
                "\n\n---\n\n".join(results)
            )

    # Fallback — tidak ada yang cocok
    all_headings = [s['heading'] for s in sections[:10]]
    return (
        f"ℹ️ Tidak ditemukan informasi relevan untuk: **\"{query}\"**\n\n"
        f"**Bagian yang tersedia dalam dokumen:**\n" +
        "\n".join(f"- {h}" for h in all_headings) +
        "\n\n_Coba gunakan kata kunci yang lebih spesifik atau tanyakan tentang bagian di atas._"
    )


# ─────────────────────────────────────────────────────────────────────────────
# CLAUDE API CHAT — diskusi isi dokumen
# ─────────────────────────────────────────────────────────────────────────────

def _build_doc_context(sections: list, max_chars: int = 14000) -> str:
    """Bangun teks konteks dari sections, potong jika terlalu panjang."""
    lines, total = [], 0
    for s in sections:
        chunk = f"\n## {s['heading']}\n" + "\n".join(s['paragraphs']) + "\n"
        if total + len(chunk) > max_chars:
            sisa = max_chars - total
            if sisa > 100:
                lines.append(chunk[:sisa] + "\n[...dokumen dipotong...]")
            break
        lines.append(chunk)
        total += len(chunk)
    return "\n".join(lines)

def _claude_chat(system: str, messages: list) -> str:
    """Kirim chat ke Z.ai API (GLM), return teks jawaban."""
    import urllib.request, json

    ZAI_API_KEY = "bd7f64d4e11642599ca8d1772e89521c.imnp62IRucfcV4bA"   # ← ganti dengan API key Z.ai kamu

    all_messages = [{"role": "system", "content": system}] + messages

    payload = json.dumps({
        "model": "glm-4-flash",
        "messages": all_messages,
        "max_tokens": 1024,
        "temperature": 0.7,
    }).encode()

    req = urllib.request.Request(
        "https://api.z.ai/api/paas/v4/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {ZAI_API_KEY}",
        },
        method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            result = json.loads(r.read())
            return result['choices'][0]['message']['content']
    except Exception as e:
        return f"❌ Z.ai Error: {e}"

# ── TAMPILAN CHAT — selalu tampil di dashboard ────────────────────────────────

# Parse dokumen — selalu refresh jika ada file hasil tapi sections masih kosong
_src = st.session_state.get('_final_tr_file') or st.session_state.get('_final_opt_file')
_cached_sections = st.session_state.get('_doc_sections', [])

if _src and os.path.exists(_src) and len(_cached_sections) == 0:
    # Ada file baru tapi belum diparsing — parse sekarang
    st.session_state['_doc_sections'] = _parse_doc_structure(_src)
elif '_doc_sections' not in st.session_state:
    st.session_state['_doc_sections'] = []

sections = st.session_state.get('_doc_sections', [])
n_sec    = len(sections)
n_para   = sum(len(s['paragraphs']) for s in sections)
has_doc  = n_sec > 0

st.divider()

# ── Asisten AI — smooth expand/collapse via st.expander ───────────────────
import streamlit.components.v1 as _components

_exp_label = (
    f"🤖 Asisten Dokumen AI  ·  {n_sec} bagian  ·  {n_para} paragraf"
    if has_doc else
    "🤖 Asisten Dokumen AI  ·  Tanya seputar ISO/SNI"
)

# Inject CSS khusus expander AI tepat sebelum render
st.markdown("""
<style>
/* Wrapper expander */
div[data-testid="stExpander"] {
    border: 1.5px solid rgba(99,102,241,0.6) !important;
    border-radius: 14px !important;
    overflow: hidden !important;
    box-shadow: 0 4px 24px rgba(99,102,241,0.3) !important;
    background: rgba(20,18,50,0.55) !important;
}
/* Header tombol */
div[data-testid="stExpander"] summary {
    background: linear-gradient(135deg, #6366f1 0%, #4f46e5 50%, #4338ca 100%) !important;
    padding: 0.82rem 1.2rem !important;
    border-radius: 12px !important;
    box-shadow: 0 3px 14px rgba(99,102,241,0.45) !important;
}
div[data-testid="stExpander"] summary:hover {
    background: linear-gradient(135deg, #818cf8 0%, #6366f1 50%, #4f46e5 100%) !important;
    box-shadow: 0 6px 22px rgba(99,102,241,0.6) !important;
}
/* Teks dalam header */
div[data-testid="stExpander"] summary span,
div[data-testid="stExpander"] summary p,
div[data-testid="stExpander"] summary div {
    color: #fff !important;
    font-weight: 700 !important;
}
/* Ikon panah */
div[data-testid="stExpander"] summary svg {
    stroke: #fff !important;
    color: #fff !important;
}
</style>
""", unsafe_allow_html=True)

with st.expander(_exp_label, expanded=True):

    # Badge status
    doc_badge = (
        f"<span style='background:rgba(16,185,129,0.12);border:1px solid rgba(16,185,129,0.3);"
        f"color:#6ee7b7;font-size:0.75rem;padding:0.2rem 0.7rem;border-radius:99px;'>"
        f"📑 {n_sec} bagian · {n_para} paragraf</span>"
        if has_doc else
        f"<span style='background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.1);"
        f"color:rgba(255,255,255,0.3);font-size:0.75rem;padding:0.2rem 0.7rem;border-radius:99px;'>"
        f"📄 Belum ada dokumen — jawab hal umum ISO/SNI</span>"
    )
    st.markdown(
        f"<div style='display:flex;gap:0.7rem;margin-bottom:0.8rem;flex-wrap:wrap;'>"
        f"{doc_badge}"
        f"<span style='background:rgba(99,102,241,0.1);border:1px solid rgba(99,102,241,0.25);"
        f"color:#a5b4fc;font-size:0.75rem;padding:0.2rem 0.7rem;border-radius:99px;'>"
        f"✨ Z.ai GLM</span></div>",
        unsafe_allow_html=True
    )

    # ── Pemilih Suara TTS ────────────────────────────────────────────────────
    _components.html("""
<div id="tts-voice-bar" style="margin-bottom:8px;display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
  <span style="color:rgba(165,180,252,0.7);font-size:0.73rem;font-family:sans-serif;">🎙️ Suara:</span>
  <select id="tts-voice-select"
    style="background:rgba(15,23,42,0.85);border:1.5px solid rgba(99,102,241,0.35);
           color:#c7d2fe;font-size:0.73rem;padding:3px 8px;border-radius:8px;
           font-family:sans-serif;cursor:pointer;outline:none;max-width:260px;">
    <option value="">⏳ Memuat daftar suara...</option>
  </select>
  <select id="tts-rate-select"
    style="background:rgba(15,23,42,0.85);border:1.5px solid rgba(99,102,241,0.35);
           color:#c7d2fe;font-size:0.73rem;padding:3px 8px;border-radius:8px;
           font-family:sans-serif;cursor:pointer;outline:none;">
    <option value="0.7">🐢 Lambat</option>
    <option value="0.92" selected>▶️ Normal</option>
    <option value="1.2">⚡ Cepat</option>
    <option value="1.5">🚀 Sangat Cepat</option>
  </select>
</div>
<script>
(function(){
  var sel = document.getElementById('tts-voice-select');

  function populateVoices(){
    var voices = window.speechSynthesis.getVoices();
    if(!voices.length) return;
    sel.innerHTML = '';

    var id_voices = voices.filter(function(v){
      return v.lang.startsWith('id') || v.lang.startsWith('ms');
    });
    var en_voices = voices.filter(function(v){
      return v.lang.startsWith('en');
    }).slice(0, 10);

    if(id_voices.length){
      var og = document.createElement('optgroup');
      og.label = '🇮🇩 Indonesia / Melayu';
      id_voices.forEach(function(v){
        var o = document.createElement('option');
        o.value = v.name;
        o.textContent = v.name + ' (' + v.lang + ')';
        og.appendChild(o);
      });
      sel.appendChild(og);
    }

    if(en_voices.length){
      var og2 = document.createElement('optgroup');
      og2.label = '🌐 English (fallback)';
      en_voices.forEach(function(v){
        var o = document.createElement('option');
        o.value = v.name;
        o.textContent = v.name + ' (' + v.lang + ')';
        og2.appendChild(o);
      });
      sel.appendChild(og2);
    }

    sel.onchange = function(){
      try{ localStorage.setItem('tts_voice', sel.value); }catch(e){}
    };
    try{
      var saved = localStorage.getItem('tts_voice');
      if(saved && sel.querySelector('option[value="'+saved+'"]')) sel.value = saved;
    }catch(e){}
  }

  if(window.speechSynthesis.getVoices().length > 0){
    populateVoices();
  } else {
    window.speechSynthesis.onvoiceschanged = populateVoices;
    setTimeout(function(){ populateVoices(); }, 500);
  }
})();
</script>
""", height=52)

    # TTS helper — teks disimpan di hidden element, bukan inline onclick
    def _tts_html(text: str, btn_id: str) -> str:
        import base64
        b64 = base64.b64encode(text[:3000].encode('utf-8')).decode('ascii')
        return f"""
<div style="margin-top:5px;display:flex;gap:6px;flex-wrap:wrap;">
  <span id="tts_data_{btn_id}" style="display:none">{b64}</span>
  <button id="btn_speak_{btn_id}"
    onclick="(function(){{
      if(!('speechSynthesis' in window)){{alert('Browser tidak mendukung TTS');return;}}
      window.speechSynthesis.cancel();

      var raw = document.getElementById('tts_data_{btn_id}').textContent;
      var txt = decodeURIComponent(escape(atob(raw)));
      var btn = document.getElementById('btn_speak_{btn_id}');

      // Ambil pilihan suara & kecepatan dari selector di frame induk
      var voiceName = '';
      var rate = 0.92;
      try {{
        var topDoc = window.top.document;
        var vSel = topDoc.getElementById('tts-voice-select');
        var rSel = topDoc.getElementById('tts-rate-select');
        if(vSel) voiceName = vSel.value;
        if(rSel) rate = parseFloat(rSel.value) || 0.92;
      }} catch(e) {{
        // fallback: coba localStorage
        try{{ voiceName = localStorage.getItem('tts_voice') || ''; }}catch(e2){{}}
      }}

      var voices = window.speechSynthesis.getVoices();
      var chosenVoice = null;
      if(voiceName) chosenVoice = voices.find(function(v){{return v.name===voiceName;}});
      if(!chosenVoice) chosenVoice = voices.find(function(v){{return v.lang.startsWith('id');}});
      if(!chosenVoice) chosenVoice = voices.find(function(v){{return v.lang.startsWith('ms');}});
      if(!chosenVoice) chosenVoice = voices.find(function(v){{return v.lang.startsWith('en');}});

      // Pecah per kalimat agar tidak terpotong browser
      var sentences = txt.match(/[^.!?\\n]{{1,220}}[.!?\\n]?/g) || [txt];
      btn.textContent = '⏳ Membaca...';
      btn.style.borderColor = 'rgba(16,185,129,0.6)';
      btn.style.color = '#6ee7b7';

      function speakChunk(i){{
        if(i >= sentences.length){{
          btn.textContent='🔊 Bacakan';
          btn.style.borderColor='rgba(99,102,241,0.4)';
          btn.style.color='#a5b4fc';
          return;
        }}
        var u = new SpeechSynthesisUtterance(sentences[i]);
        u.lang = 'id-ID';
        u.rate = rate;
        u.pitch = 1.0;
        if(chosenVoice) u.voice = chosenVoice;
        u.onend  = function(){{ speakChunk(i+1); }};
        u.onerror = function(){{
          btn.textContent='🔊 Bacakan';
          btn.style.borderColor='rgba(99,102,241,0.4)';
          btn.style.color='#a5b4fc';
        }};
        window.speechSynthesis.speak(u);
      }}

      function startSpeak(){{
        if(window.speechSynthesis.getVoices().length===0){{
          window.speechSynthesis.onvoiceschanged=function(){{ speakChunk(0); }};
        }} else {{
          speakChunk(0);
        }}
      }}
      startSpeak();
    }})()"
    style="background:rgba(99,102,241,0.15);border:1px solid rgba(99,102,241,0.4);
           color:#a5b4fc;font-size:0.72rem;padding:4px 12px;border-radius:99px;
           cursor:pointer;font-family:sans-serif;transition:all 0.2s;">
    🔊 Bacakan
  </button>
  <button onclick="(function(){{
      window.speechSynthesis.cancel();
      var b=document.getElementById('btn_speak_{btn_id}');
      if(b){{b.textContent='🔊 Bacakan';b.style.borderColor='rgba(99,102,241,0.4)';b.style.color='#a5b4fc';}}
    }})()"
    style="background:rgba(239,68,68,0.1);border:1px solid rgba(239,68,68,0.3);
           color:#fca5a5;font-size:0.72rem;padding:4px 12px;border-radius:99px;
           cursor:pointer;font-family:sans-serif;transition:all 0.2s;">
    ⏹ Stop
  </button>
</div>
"""

    # Inisialisasi riwayat
    if '_chat_history' not in st.session_state:
        st.session_state['_chat_history'] = []

    # Tampilkan riwayat
    for idx, msg in enumerate(st.session_state['_chat_history']):
        with st.chat_message(msg['role'], avatar='🧑' if msg['role'] == 'user' else '🤖'):
            st.markdown(msg['content'])
            if msg['role'] == 'assistant':
                _clean = re.sub(r'[*_`#>\-]+', '', msg['content'])
                _clean = re.sub(r'\s+', ' ', _clean).strip()
                _components.html(_tts_html(_clean, f"hist_{idx}"), height=44)

    # ── Voice Input — fixed position, kompatibel Streamlit Cloud ────────────
    _components.html("""
<style>
#mic-fixed-btn {
  position: fixed;
  bottom: 18px;
  right: calc(50% - 360px);   /* sejajar kanan chat area max-width 780px */
  z-index: 99999;
  background: linear-gradient(135deg, #6366f1, #4f46e5);
  border: none; color: #fff;
  font-size: 0.78rem; font-weight: 700;
  padding: 9px 16px; border-radius: 99px; cursor: pointer;
  font-family: 'Outfit', sans-serif;
  box-shadow: 0 4px 16px rgba(99,102,241,0.5);
  transition: all 0.2s; white-space: nowrap;
  display: flex; align-items: center; gap: 5px;
}
#mic-fixed-btn:hover {
  background: linear-gradient(135deg,#818cf8,#6366f1);
  box-shadow: 0 6px 22px rgba(99,102,241,0.65);
  transform: translateY(-1px);
}
#mic-fixed-btn.listening {
  background: linear-gradient(135deg,#ef4444,#dc2626) !important;
  animation: mic-pulse 1s infinite;
}
@keyframes mic-pulse {
  0%   { box-shadow: 0 0 0 0 rgba(239,68,68,0.6); }
  70%  { box-shadow: 0 0 0 10px rgba(239,68,68,0); }
  100% { box-shadow: 0 0 0 0 rgba(239,68,68,0); }
}
#mic-toast {
  display: none;
  position: fixed; bottom: 65px; left: 50%;
  transform: translateX(-50%);
  background: rgba(15,23,42,0.96);
  border: 1px solid rgba(99,102,241,0.4);
  border-radius: 12px; padding: 7px 16px;
  font-size: 0.78rem; color: #c7d2fe;
  font-family: sans-serif; z-index: 99999;
  box-shadow: 0 4px 20px rgba(0,0,0,0.4);
  white-space: nowrap; pointer-events: none;
  max-width: 90vw; overflow: hidden;
  text-overflow: ellipsis;
}
/* Responsif: geser tombol ke kiri saat layar kecil */
@media (max-width: 860px) {
  #mic-fixed-btn { right: 12px; }
}
</style>

<button id="mic-fixed-btn" onclick="toggleMic()">🎤 Bicara</button>
<div id="mic-toast"></div>

<script>
(function(){
  // ── State tersimpan di window agar bertahan antar re-render iframe ────────
  if (!window._mic) window._mic = { rec: null, listening: false, timer: null };
  var M = window._mic;

  var btn   = document.getElementById('mic-fixed-btn');
  var toast = document.getElementById('mic-toast');

  function showToast(msg, ms) {
    toast.textContent = msg;
    toast.style.display = 'block';
    clearTimeout(M.timer);
    if (ms) M.timer = setTimeout(function(){ toast.style.display='none'; }, ms);
  }

  function resetBtn() {
    btn.innerHTML = '🎤 Bicara';
    btn.classList.remove('listening');
  }

  function forceStop() {
    if (M.rec) { try { M.rec.abort(); } catch(e){} M.rec = null; }
    M.listening = false;
    resetBtn();
  }

  function sendToChat(txt) {
    /* Coba inject ke parent Streamlit — works lokal.
       Di Cloud, fallback ke clipboard + notif manual. */
    var sent = false;

    // Coba window, window.parent, window.top secara berturut
    var targets = [];
    try { targets.push(window); } catch(e){}
    try { if (window.parent !== window) targets.push(window.parent); } catch(e){}
    try { if (window.top !== window && window.top !== window.parent) targets.push(window.top); } catch(e){}

    for (var i = 0; i < targets.length && !sent; i++) {
      try {
        var w = targets[i];
        var ta = w.document.querySelector('div[data-testid="stChatInput"] textarea');
        if (!ta) continue;

        var setter = Object.getOwnPropertyDescriptor(w.HTMLTextAreaElement.prototype, 'value').set;
        setter.call(ta, txt);
        ta.dispatchEvent(new Event('input',  {bubbles:true}));
        ta.dispatchEvent(new Event('change', {bubbles:true}));

        setTimeout(function(){
          ['keydown','keypress','keyup'].forEach(function(evName){
            ta.dispatchEvent(new w.KeyboardEvent(evName, {
              key:'Enter', code:'Enter', keyCode:13, which:13,
              bubbles:true, cancelable:true
            }));
          });
          toast.style.display = 'none';
        }, 150);

        sent = true;
      } catch(e) {}
    }

    if (!sent) {
      // Fallback: copy ke clipboard, minta user paste manual
      try {
        navigator.clipboard.writeText(txt).then(function(){
          showToast('📋 Tersalin — paste (Ctrl+V) ke kotak chat', 4000);
        });
      } catch(e) {
        showToast('💬 ' + txt.substring(0,60) + (txt.length>60?'...':''), 5000);
      }
    }
  }

  window.toggleMic = function() {
    if (M.listening) {
      forceStop();
      showToast('⏹ Dihentikan', 1500);
      return;
    }

    var SR = window.SpeechRecognition || window.webkitSpeechRecognition;
    if (!SR) {
      showToast('❌ Gunakan Chrome/Edge untuk voice input', 3000);
      return;
    }

    forceStop();

    var rec = new SR();
    rec.lang = 'id-ID';
    rec.continuous = false;
    rec.interimResults = true;
    rec.maxAlternatives = 1;
    M.rec = rec;

    var lastTxt = '';

    rec.onstart = function() {
      M.listening = true;
      btn.innerHTML = '🔴 Stop';
      btn.classList.add('listening');
      showToast('🎤 Sedang mendengarkan...');
    };

    rec.onresult = function(e) {
      var interim='', final_t='';
      for (var i=e.resultIndex; i<e.results.length; i++) {
        if (e.results[i].isFinal) final_t += e.results[i][0].transcript;
        else interim += e.results[i][0].transcript;
      }
      lastTxt = final_t || interim;
      showToast('💬 ' + lastTxt);
    };

    rec.onend = function() {
      M.rec = null; M.listening = false;
      resetBtn();
      var txt = lastTxt.trim(); lastTxt = '';
      if (!txt) { showToast('⚠️ Tidak terdeteksi — coba lagi', 2500); return; }
      sendToChat(txt);
    };

    rec.onerror = function(e) {
      M.rec = null; M.listening = false; resetBtn();
      if (e.error === 'aborted') return;
      var msgs = {
        'no-speech'    : '⚠️ Tidak ada suara — coba lagi',
        'not-allowed'  : '❌ Izin mikrofon ditolak di browser',
        'audio-capture': '❌ Mikrofon tidak ditemukan',
        'network'      : '❌ Gangguan jaringan'
      };
      showToast(msgs[e.error] || ('❌ Error: ' + e.error), 3000);
    };

    rec.start();
  };
})();
</script>
""", height=56)

    # Input teks
    user_input = st.chat_input(
        "Tanya seputar ISO/SNI atau isi dokumen...?"
    )

    if user_input:
        st.session_state['_chat_history'].append({'role': 'user', 'content': user_input})
        with st.chat_message('user', avatar='🧑'):
            st.markdown(user_input)

        with st.chat_message('assistant', avatar='🤖'):
            with st.spinner("asistant sedang menganalisis..."):
                if has_doc:
                    doc_ctx = _build_doc_context(sections, max_chars=14000)
                    system_prompt = (
                        "Kamu adalah asisten ahli standar ISO/SNI yang membantu pengguna memahami dokumen. "
                        "Jawab HANYA berdasarkan isi dokumen berikut. "
                        "Gunakan Bahasa Indonesia yang jelas, terstruktur, dan akurat. "
                        "Jika informasi tidak ada dalam dokumen, katakan dengan jujur.\n\n"
                        f"=== ISI DOKUMEN ===\n{doc_ctx}\n==================="
                    )
                else:
                    system_prompt = (
                        "Kamu adalah asisten ahli standar ISO/SNI dan dokumen teknis BSN. "
                        "Jawab dalam Bahasa Indonesia dengan jelas dan akurat. "
                        "Belum ada dokumen — jawab berdasarkan pengetahuan umum ISO, SNI, IEC, dan standardisasi."
                    )
                api_messages = [
                    {"role": m['role'], "content": m['content']}
                    for m in st.session_state['_chat_history']
                ]
                reply = _claude_chat(system_prompt, api_messages)

            st.markdown(reply)
            _clean = re.sub(r'[*_`#>\-]+', '', reply)
            _clean = re.sub(r'\s+', ' ', _clean).strip()
            _components.html(_tts_html(_clean, f"new_{int(time.time())}"), height=44)
            st.session_state['_chat_history'].append({'role': 'assistant', 'content': reply})

    if st.session_state.get('_chat_history'):
        if st.button("🗑️ Hapus Riwayat", key="clear_chat"):
            st.session_state['_chat_history'] = []
            st.rerun()

# --- FOOTER ---
_now = time.strftime("%H:%M")
st.markdown(
    f"<div class='footer'>"
    f"© 2026 <span>ISO/SNI Master Formating</span> · Transportasi dan Teknologi Informasi · All rights reserved."
    f"<br><span style='font-size:0.68rem;opacity:0.5;'>🛡️ File temporer dihapus otomatis setiap 30 menit &nbsp;·&nbsp; Terakhir dicek: {_now}</span>"
    f"</div>",
    unsafe_allow_html=True
)
